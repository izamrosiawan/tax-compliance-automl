import re
import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def sanitize_math(text: str) -> str:
    """Translates LaTeX math notation into clean, standard Unicode symbols."""
    t = text
    t = t.replace(r'\delta_{\text{cash}}', 'δ_cash')
    t = t.replace(r'\delta_{\text{inv}}', 'δ_inv')
    t = t.replace(r'\delta_{\text{log}}', 'δ_log')
    t = t.replace(r'\delta', 'δ')
    t = t.replace(r'\Delta', 'Δ')
    t = t.replace(r'\varepsilon', 'ε')
    t = t.replace(r'\approx', '≈')
    t = t.replace(r'\times', 'x')
    t = t.replace(r'\pm', '±')
    t = t.replace(r'\le', '≤')
    t = t.replace(r'\ge', '≥')
    t = t.replace(r'\in', '∈')
    t = t.replace(r'\sim', '~')
    t = t.replace(r'\mathbb{I}', 'I')
    t = t.replace(r'\mathcal{N}', 'N')
    t = t.replace(r'\%', '%')
    t = t.replace(r'\_', '_')
    t = t.replace(r'\beta', 'β')
    t = t.replace(r'\lambda', 'λ')
    t = t.replace(r'\mu', 'μ')
    t = t.replace(r'\sigma', 'σ')
    t = t.replace(r'\phi', 'ϕ')
    t = t.replace(r'\cdot', '·')
    t = t.replace(r'\left(', '(')
    t = t.replace(r'\right)', ')')
    t = t.replace(r'\quad', ' ')
    t = re.sub(r'\\frac\{([^}]*)\}\{([^}]*)\}', r'(\1/\2)', t)
    t = re.sub(r'\\text\{([^}]*)\}', r'\1', t)
    t = t.replace('$$', '')
    t = t.replace('$', '')
    t = t.replace('\\', '')
    return t

def add_formatted_runs(paragraph, text: str, default_font_size=Pt(11), is_italic_block=False):
    """
    Renders text by converting Markdown asterisks (*italic* and **bold**) 
    into true Word Runs with zero literal asterisks.
    """
    clean = sanitize_math(text)
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*)', clean)
    
    for token in tokens:
        if not token:
            continue
        if token.startswith('**') and token.endswith('**') and len(token) >= 4:
            inner = token[2:-2]
            r = paragraph.add_run(inner)
            r.bold = True
            r.italic = is_italic_block
            r.font.name = 'Times New Roman'
            r.font.size = default_font_size
        elif token.startswith('*') and token.endswith('*') and len(token) >= 2:
            inner = token[1:-1]
            r = paragraph.add_run(inner)
            r.italic = True
            r.font.name = 'Times New Roman'
            r.font.size = default_font_size
        else:
            r = paragraph.add_run(token)
            r.italic = is_italic_block
            r.font.name = 'Times New Roman'
            r.font.size = default_font_size

def set_cell_borders(cell, is_header=False, is_last=False):
    """Professional Academic Three-Line Table formatting."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(r'<w:tcBorders {}/>'.format(nsdecls('w')))
    
    if is_header:
        top_el = parse_xml(f'<w:top {nsdecls("w")} w:val="single" w:sz="12" w:space="0" w:color="000000"/>')
        bot_el = parse_xml(f'<w:bottom {nsdecls("w")} w:val="single" w:sz="6" w:space="0" w:color="000000"/>')
    elif is_last:
        top_el = parse_xml(f'<w:top {nsdecls("w")} w:val="none"/>')
        bot_el = parse_xml(f'<w:bottom {nsdecls("w")} w:val="single" w:sz="12" w:space="0" w:color="000000"/>')
    else:
        top_el = parse_xml(f'<w:top {nsdecls("w")} w:val="none"/>')
        bot_el = parse_xml(f'<w:bottom {nsdecls("w")} w:val="none"/>')
        
    tcBorders.append(top_el)
    tcBorders.append(bot_el)
    tcBorders.append(parse_xml(f'<w:left {nsdecls("w")} w:val="none"/>'))
    tcBorders.append(parse_xml(f'<w:right {nsdecls("w")} w:val="none"/>'))
    tcPr.append(tcBorders)

def render_table(doc, tbl_lines):
    if len(tbl_lines) < 2:
        return
    rows_data = []
    for l in tbl_lines:
        if re.match(r'^\s*\|?\s*:?---', l):
            continue
        cells = [c.strip() for c in l.strip().strip('|').split('|')]
        if any(cells):
            rows_data.append(cells)
    if not rows_data:
        return
    
    cols = max(len(r) for r in rows_data)
    tbl = doc.add_table(rows=len(rows_data), cols=cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set explicit table width to 6.3 inches (full text width)
    tbl_pr = tbl._tbl.tblPr
    tbl_w = parse_xml(r'<w:tblW {} w:w="9072" w:type="dxa"/>'.format(nsdecls('w')))
    tbl_pr.append(tbl_w)
    
    # Distribute column widths intelligently
    total_width_inches = 6.3
    # If 9 columns (Table 3):
    if cols == 9:
        col_widths = [1.2, 0.6, 0.8, 0.7, 0.5, 0.5, 0.6, 0.5, 0.9]
    elif cols == 5:
        # Table 2 or Table 4
        col_widths = [1.2, 0.9, 1.4, 2.0, 0.8]
    elif cols == 4:
        col_widths = [1.8, 1.8, 1.35, 1.35]
    elif cols == 7:
        col_widths = [0.9, 0.9, 0.9, 0.9, 0.9, 0.9, 0.9]
    else:
        col_widths = [total_width_inches / cols] * cols
        
    for i, row in enumerate(rows_data):
        is_header = (i == 0)
        is_last = (i == len(rows_data) - 1)
        for j, cell_text in enumerate(row):
            if j >= cols:
                continue
            cell = tbl.cell(i, j)
            cell.width = Inches(col_widths[min(j, len(col_widths)-1)])
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.10
            
            clean_str = sanitize_math(cell_text).replace('*', '')
            if j == 0 and not is_header and len(clean_str) > 6:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif any(ch.isdigit() for ch in clean_str) and len(clean_str) < 30:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            # Table font size standard: 9.5 pt for readable text, 8.5 pt for wide 9-col table
            table_font_pt = Pt(8.5) if cols >= 8 else Pt(9.5)
            add_formatted_runs(p, cell_text, default_font_size=table_font_pt)
            set_cell_borders(cell, is_header=is_header, is_last=is_last)
            if is_header:
                shd = parse_xml(r'<w:shd {} w:fill="F5F5F5"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shd)
                
    p_sp = doc.add_paragraph()
    p_sp.paragraph_format.space_after = Pt(6)

def generate_sinta_docx():
    doc = docx.Document()
    
    # 1. Page Margins (Standard SINTA 2: 2.54 cm / 1.0 inch all sides)
    for sec in doc.sections:
        sec.page_width = Inches(8.27)
        sec.page_height = Inches(11.69)
        sec.top_margin = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin = Inches(1.0)
        sec.right_margin = Inches(1.0)
        
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Times New Roman'
    font.size = Pt(11) # Standard journal body font
    font.color.rgb = RGBColor(0, 0, 0)
    
    with open("paper.md", "r", encoding="utf-8") as f:
        md_content = f.read()

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(6)
    p_title.paragraph_format.line_spacing = 1.15
    r_t = p_title.add_run("Kerangka Benchmark Simulasi Non-Sirkular untuk Pemeringkatan Risiko Kepatuhan Pajak Pedagang Daring melalui Integrasi Data Transaksi Gerbang Pembayaran, Logistik, dan Indikator Regional Berbasis Data BPS")
    r_t.bold = True
    r_t.font.name = 'Times New Roman'
    r_t.font.size = Pt(14)
    
    # English Title
    p_etitle = doc.add_paragraph()
    p_etitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_etitle.paragraph_format.space_after = Pt(14)
    r_et = p_etitle.add_run("(A Non-Circular Synthetic Simulation Benchmark for E-Commerce Merchant Tax Compliance Risk Ranking Integrating Digital Payment Gateway, Logistics, and BPS Regional Indicators)")
    r_et.italic = True
    r_et.font.name = 'Times New Roman'
    r_et.font.size = Pt(11)
    
    # Authors
    p_auth = doc.add_paragraph()
    p_auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_auth.paragraph_format.space_after = Pt(2)
    r1 = p_auth.add_run("Izam Rosiawan")
    r1.bold = True
    r1.font.size = Pt(11)
    r1_sup = p_auth.add_run("1*")
    r1_sup.font.superscript = True
    p_auth.add_run(", ")
    r2 = p_auth.add_run("Sulthan")
    r2.bold = True
    r2.font.size = Pt(11)
    r2_sup = p_auth.add_run("2")
    r2_sup.font.superscript = True
    
    # Affiliation
    p_aff = doc.add_paragraph()
    p_aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_aff.paragraph_format.space_after = Pt(2)
    p_aff.paragraph_format.line_spacing = 1.05
    r_aff1 = p_aff.add_run("1Program Studi Sains Data, Fakultas Informatika, Telkom University Kampus Surabaya, Indonesia\n")
    r_aff1.font.size = Pt(10)
    r_aff2 = p_aff.add_run("2Direktorat Kampus Surabaya, Telkom University Kampus Surabaya, Indonesia\n")
    r_aff2.font.size = Pt(10)
    r_aff3 = p_aff.add_run("*Penulis Korespondensi: izamrosiawan@student.telkomuniversity.ac.id")
    r_aff3.font.size = Pt(9.5)
    r_aff3.italic = True
    p_aff.paragraph_format.space_after = Pt(12)
    
    # Top Divider Line
    p_line1 = doc.add_paragraph()
    p_line1.paragraph_format.space_after = Pt(6)
    p_border1 = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="6" w:space="1" w:color="000000"/></w:pBdr>')
    p_line1._p.get_or_add_pPr().append(p_border1)
    
    # Abstrak Indo
    p_abs = doc.add_paragraph()
    p_abs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_abs.paragraph_format.space_after = Pt(4)
    p_abs.paragraph_format.line_spacing = 1.0
    r_h_id = p_abs.add_run("ABSTRAK — ")
    r_h_id.bold = True
    r_h_id.font.size = Pt(10)
    
    m_abs_id = re.search(r'### ABSTRAK\s*\n(.*?)\n\s*\*\*Kata Kunci:\*\*\s*(.*?)\n', md_content, re.DOTALL)
    abs_id_txt = m_abs_id.group(1).strip() if m_abs_id else ""
    kw_id_txt = m_abs_id.group(2).strip() if m_abs_id else ""
    add_formatted_runs(p_abs, abs_id_txt, default_font_size=Pt(10))
    
    p_kw_id = doc.add_paragraph()
    p_kw_id.paragraph_format.space_after = Pt(8)
    r_kwt_id = p_kw_id.add_run("Kata Kunci: ")
    r_kwt_id.bold = True
    r_kwt_id.font.size = Pt(10)
    r_kw_body_id = p_kw_id.add_run(sanitize_math(kw_id_txt))
    r_kw_body_id.italic = True
    r_kw_body_id.font.size = Pt(10)
    
    # Abstrak Inggris
    p_eabs = doc.add_paragraph()
    p_eabs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_eabs.paragraph_format.space_after = Pt(4)
    p_eabs.paragraph_format.line_spacing = 1.0
    r_h_en = p_eabs.add_run("ABSTRACT — ")
    r_h_en.bold = True
    r_h_en.font.size = Pt(10)
    
    m_abs_en = re.search(r'### ABSTRACT\s*\n\*(.*?)\*\n\s*\*\*Keywords:\*\*\s*(.*?)\n', md_content, re.DOTALL)
    abs_en_txt = m_abs_en.group(1).strip() if m_abs_en else ""
    kw_en_txt = m_abs_en.group(2).strip() if m_abs_en else ""
    add_formatted_runs(p_eabs, abs_en_txt, default_font_size=Pt(10), is_italic_block=True)
    
    p_kw_en = doc.add_paragraph()
    p_kw_en.paragraph_format.space_after = Pt(8)
    r_kwt_en = p_kw_en.add_run("Keywords: ")
    r_kwt_en.bold = True
    r_kwt_en.font.size = Pt(10)
    r_kw_body_en = p_kw_en.add_run(sanitize_math(kw_en_txt))
    r_kw_body_en.italic = True
    r_kw_body_en.font.size = Pt(10)
    
    # Bottom Divider Line
    p_line2 = doc.add_paragraph()
    p_line2.paragraph_format.space_after = Pt(14)
    p_border2 = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="6" w:space="1" w:color="000000"/></w:pBdr>')
    p_line2._p.get_or_add_pPr().append(p_border2)

    # Body Parsing
    body_start = md_content.find("## 1. PENDAHULUAN")
    body_text = md_content[body_start:]
    lines = body_text.split("\n")
    
    table_lines = []
    idx = 0
    while idx < len(lines):
        line = lines[idx].strip()
        
        # Table collector
        if line.startswith("|") and line.endswith("|"):
            table_lines.append(line)
            idx += 1
            continue
        elif table_lines:
            render_table(doc, table_lines)
            table_lines = []
            
        if not line:
            idx += 1
            continue
            
        # Heading 1
        if line.startswith("## "):
            h_text = sanitize_math(line[3:]).replace('*', '')
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(h_text)
            r.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)
        # Heading 2
        elif line.startswith("### "):
            h_text = sanitize_math(line[4:]).replace('*', '')
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(h_text)
            r.bold = True
            r.italic = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(11)
        # Bold-Italic Sub-headline (e.g. *Analisis Overfitting:*)
        elif line.startswith("*") and line.endswith(":*"):
            sub_text = line.strip('*').rstrip(':')
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(sub_text + ":")
            r.bold = True
            r.italic = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(11)
        # Images
        elif line.startswith("!["):
            m_img = re.search(r'!\[(.*?)\]\((.*?)\)', line)
            if m_img:
                img_path = m_img.group(2).lstrip('/')
                if os.path.exists(img_path):
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.paragraph_format.space_before = Pt(8)
                    p_img.paragraph_format.space_after = Pt(2)
                    r_img = p_img.add_run()
                    r_img.add_picture(img_path, width=Inches(6.0))
        # Image Caption
        elif line.startswith("*Gambar ") or line.startswith("Gambar "):
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_after = Pt(8)
            clean_cap = line.strip('*')
            r_cap = p_cap.add_run(clean_cap)
            r_cap.font.name = 'Times New Roman'
            r_cap.font.size = Pt(9.5)
            r_cap.italic = True
        # Table Caption
        elif line.startswith("**Tabel ") or line.startswith("Tabel "):
            p_tcap = doc.add_paragraph()
            p_tcap.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_tcap.paragraph_format.space_before = Pt(8)
            p_tcap.paragraph_format.space_after = Pt(2)
            clean_tcap = line.replace('**', '')
            r_tcap = p_tcap.add_run(clean_tcap)
            r_tcap.bold = True
            r_tcap.font.name = 'Times New Roman'
            r_tcap.font.size = Pt(10)
        # Table Footnote / Notes
        elif line.startswith("*(") and line.endswith(")*"):
            p_tnot = doc.add_paragraph()
            p_tnot.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_tnot.paragraph_format.space_after = Pt(4)
            clean_not = line.strip('*')
            r_tnot = p_tnot.add_run(clean_not)
            r_tnot.italic = True
            r_tnot.font.name = 'Times New Roman'
            r_tnot.font.size = Pt(9.0)
        # References (hanging indent)
        elif line.startswith("* "):
            p_ref = doc.add_paragraph()
            p_ref.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_ref.paragraph_format.left_indent = Inches(0.3)
            p_ref.paragraph_format.first_line_indent = Inches(-0.3)
            p_ref.paragraph_format.space_after = Pt(4)
            p_ref.paragraph_format.line_spacing = 1.10
            add_formatted_runs(p_ref, line[2:].strip(), default_font_size=Pt(10))
        # Lists (a., b., c., 1., 2., etc.)
        elif re.match(r'^[a-d]\.\s', line) or re.match(r'^\d\.\s', line):
            p_li = doc.add_paragraph()
            p_li.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_li.paragraph_format.left_indent = Inches(0.25)
            p_li.paragraph_format.space_after = Pt(3)
            p_li.paragraph_format.line_spacing = 1.15
            add_formatted_runs(p_li, line, default_font_size=Pt(11))
        # Code/Verbatim Diagram Box
        elif line.startswith("```"):
            code_lines = []
            idx += 1
            while idx < len(lines) and not lines[idx].strip().startswith("```"):
                code_lines.append(lines[idx])
                idx += 1
            p_code = doc.add_paragraph()
            p_code.paragraph_format.left_indent = Inches(0.25)
            p_code.paragraph_format.space_before = Pt(4)
            p_code.paragraph_format.space_after = Pt(6)
            r_code = p_code.add_run("\n".join(code_lines))
            r_code.font.name = 'Consolas'
            r_code.font.size = Pt(8.5)
        # Standalone equation
        elif line.startswith("$$") and line.endswith("$$"):
            p_eq = doc.add_paragraph()
            p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_eq.paragraph_format.space_before = Pt(4)
            p_eq.paragraph_format.space_after = Pt(4)
            eq_text = sanitize_math(line)
            r_eq = p_eq.add_run(eq_text)
            r_eq.font.name = 'Times New Roman'
            r_eq.italic = True
            r_eq.font.size = Pt(11)
        # Standard Paragraph
        else:
            p_p = doc.add_paragraph()
            p_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_p.paragraph_format.first_line_indent = Inches(0.35)
            p_p.paragraph_format.space_after = Pt(5)
            p_p.paragraph_format.line_spacing = 1.15
            add_formatted_runs(p_p, line, default_font_size=Pt(11))
            
        idx += 1
        
    if table_lines:
        render_table(doc, table_lines)
        
    out_path = "tax_compliance_manuscript_sinta2_final.docx"
    doc.save(out_path)
    print("Document successfully generated with standard comfortable typography and table widths.")

if __name__ == "__main__":
    generate_sinta_docx()
