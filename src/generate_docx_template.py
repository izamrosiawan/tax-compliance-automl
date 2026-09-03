import re
import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_sinta2_docx(paper_md_path: str, output_docx_path: str):
    doc = docx.Document()
    
    # Page setup - A4 Standard, Margins 2.5 cm (approx 1 inch)
    for section in doc.sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Styles setup
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Times New Roman'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0, 0, 0)
    
    with open(paper_md_path, 'r', encoding='utf-8') as f:
        text = f.read()

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(12)
    run_title = p_title.add_run("Kerangka Benchmark Simulasi Non-Sirkular untuk Pemeringkatan Risiko Kepatuhan Pajak Pedagang Daring melalui Integrasi Data Transaksi Gerbang Pembayaran, Logistik, dan Indikator Regional Berbasis Data BPS")
    run_title.bold = True
    run_title.font.size = Pt(14)
    run_title.font.name = 'Times New Roman'
    
    # English Title
    p_etitle = doc.add_paragraph()
    p_etitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_etitle.paragraph_format.space_after = Pt(14)
    run_etitle = p_etitle.add_run("(A Non-Circular Synthetic Simulation Benchmark for E-Commerce Merchant Tax Compliance Risk Ranking Integrating Digital Payment Gateway, Logistics, and BPS Regional Indicators)")
    run_etitle.italic = True
    run_etitle.font.size = Pt(11)
    run_etitle.font.name = 'Times New Roman'
    
    # Authors
    p_author = doc.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_author.paragraph_format.space_after = Pt(4)
    r_a1 = p_author.add_run("Izam Rosiawan")
    r_a1.bold = True
    r_a1.font.size = Pt(11)
    r_sup1 = p_author.add_run("1*")
    r_sup1.font.superscript = True
    r_sep = p_author.add_run(", ")
    r_a2 = p_author.add_run("Sulthan")
    r_a2.bold = True
    r_a2.font.size = Pt(11)
    r_sup2 = p_author.add_run("2")
    r_sup2.font.superscript = True
    
    # Affiliation
    p_aff = doc.add_paragraph()
    p_aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_aff.paragraph_format.space_after = Pt(2)
    r_aff1 = p_aff.add_run("1Program Studi Sains Data, Fakultas Informatika, Telkom University, Kampus Surabaya, Indonesia\n")
    r_aff1.font.size = Pt(9.5)
    r_aff2 = p_aff.add_run("2Direktorat Kampus Surabaya, Telkom University, Kampus Surabaya, Indonesia\n")
    r_aff2.font.size = Pt(9.5)
    r_aff3 = p_aff.add_run("*Penulis Korespondensi: izamrosiawan@student.telkomuniversity.ac.id")
    r_aff3.font.size = Pt(9.5)
    r_aff3.italic = True
    p_aff.paragraph_format.space_after = Pt(14)
    
    # Divider line
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(10)
    p_div_border = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/></w:pBdr>')
    p_div._p.get_or_add_pPr().append(p_div_border)
    
    # Abstracts Box
    # Indonesia
    p_abs_title = doc.add_paragraph()
    p_abs_title.paragraph_format.space_after = Pt(2)
    r_abt = p_abs_title.add_run("ABSTRAK")
    r_abt.bold = True
    r_abt.font.size = Pt(10)
    
    # Extract Abstrak Indo
    m_abs_id = re.search(r'### ABSTRAK\s*\n(.*?)\n\s*\*\*Kata Kunci:\*\*\s*(.*?)\n', text, re.DOTALL)
    if m_abs_id:
        abs_text_id = m_abs_id.group(1).strip()
        kw_id = m_abs_id.group(2).strip()
        p_abs = doc.add_paragraph()
        p_abs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_abs.paragraph_format.space_after = Pt(4)
        p_abs.paragraph_format.line_spacing = 1.0
        r_abs = p_abs.add_run(abs_text_id)
        r_abs.font.size = Pt(9.5)
        
        p_kw = doc.add_paragraph()
        p_kw.paragraph_format.space_after = Pt(12)
        r_kwt = p_kw.add_run("Kata Kunci: ")
        r_kwt.bold = True
        r_kwt.font.size = Pt(9.5)
        r_kw = p_kw.add_run(kw_id)
        r_kw.italic = True
        r_kw.font.size = Pt(9.5)
        
    # English Abstract
    p_eabs_title = doc.add_paragraph()
    p_eabs_title.paragraph_format.space_after = Pt(2)
    r_eabt = p_eabs_title.add_run("ABSTRACT")
    r_eabt.bold = True
    r_eabt.font.size = Pt(10)
    
    m_abs_en = re.search(r'### ABSTRACT\s*\n\*(.*?)\*\n\s*\*\*Keywords:\*\*\s*(.*?)\n', text, re.DOTALL)
    if m_abs_en:
        abs_text_en = m_abs_en.group(1).strip()
        kw_en = m_abs_en.group(2).strip()
        p_eabs = doc.add_paragraph()
        p_eabs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_eabs.paragraph_format.space_after = Pt(4)
        p_eabs.paragraph_format.line_spacing = 1.0
        r_eabs = p_eabs.add_run(abs_text_en)
        r_eabs.italic = True
        r_eabs.font.size = Pt(9.5)
        
        p_ekw = doc.add_paragraph()
        p_ekw.paragraph_format.space_after = Pt(14)
        r_ekwt = p_ekw.add_run("Keywords: ")
        r_ekwt.bold = True
        r_ekwt.font.size = Pt(9.5)
        r_ekw = p_ekw.add_run(kw_en)
        r_ekw.italic = True
        r_ekw.font.size = Pt(9.5)

    # Divider line
    p_div2 = doc.add_paragraph()
    p_div2.paragraph_format.space_after = Pt(14)
    p_div2_border = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/></w:pBdr>')
    p_div2._p.get_or_add_pPr().append(p_div2_border)

    # Body parser
    # Split text into sections starting from ## 1. PENDAHULUAN
    start_body = text.find("## 1. PENDAHULUAN")
    body_text = text[start_body:]
    
    lines = body_text.split("\n")
    in_table = False
    table_rows = []
    
    def process_table(rows):
        if not rows or len(rows) < 2:
            return
        # Parse markdown table
        parsed_rows = []
        for r in rows:
            if re.match(r'^\s*\|?\s*:?---', r):
                continue
            cells = [c.strip() for c in r.strip().strip('|').split('|')]
            if cells and any(cells):
                parsed_rows.append(cells)
        if not parsed_rows:
            return
        
        cols_count = max(len(r) for r in parsed_rows)
        tbl = doc.add_table(rows=len(parsed_rows), cols=cols_count)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = True
        
        for i, row in enumerate(parsed_rows):
            for j, cell_text in enumerate(row):
                if j < cols_count:
                    c = tbl.cell(i, j)
                    c.text = cell_text
                    # formatting
                    for p in c.paragraphs:
                        p.paragraph_format.space_before = Pt(2)
                        p.paragraph_format.space_after = Pt(2)
                        for r in p.runs:
                            r.font.name = 'Times New Roman'
                            r.font.size = Pt(8.5)
                            if i == 0:
                                r.bold = True
                    # Shading for header
                    if i == 0:
                        shd = parse_xml(r'<w:shd {} w:fill="F0F0F0"/>'.format(nsdecls('w')))
                        c._tc.get_or_add_tcPr().append(shd)
                        
        p_sp = doc.add_paragraph()
        p_sp.paragraph_format.space_after = Pt(6)

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Check table
        if line.startswith("|") and line.endswith("|"):
            table_rows.append(line)
            i += 1
            continue
        elif table_rows:
            process_table(table_rows)
            table_rows = []

        if not line:
            i += 1
            continue
            
        # Headings
        if line.startswith("## "):
            h_text = line[3:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(h_text)
            r.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(11.5)
        elif line.startswith("### "):
            h_text = line[4:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(h_text)
            r.bold = True
            r.italic = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10.5)
        elif line.startswith("!["):
            # Image handling
            m_img = re.search(r'!\[(.*?)\]\((.*?)\)', line)
            if m_img:
                img_caption = m_img.group(1)
                img_rel = m_img.group(2).lstrip('/')
                img_path = os.path.join(os.path.dirname(paper_md_path), img_rel)
                if os.path.exists(img_path):
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.paragraph_format.space_before = Pt(8)
                    p_img.paragraph_format.space_after = Pt(2)
                    run_img = p_img.add_run()
                    run_img.add_picture(img_path, width=Inches(5.5))
        elif line.startswith("*Gambar ") or line.startswith("Gambar "):
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_after = Pt(8)
            r = p_cap.add_run(line.strip('*'))
            r.italic = True
            r.font.size = Pt(9)
        elif line.startswith("**Tabel ") or line.startswith("Tabel "):
            p_tcap = doc.add_paragraph()
            p_tcap.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_tcap.paragraph_format.space_before = Pt(8)
            p_tcap.paragraph_format.space_after = Pt(2)
            r = p_tcap.add_run(line.strip('*'))
            r.bold = True
            r.font.size = Pt(9.5)
        elif line.startswith("* "):
            p_ref = doc.add_paragraph()
            p_ref.paragraph_format.left_indent = Inches(0.25)
            p_ref.paragraph_format.first_line_indent = Inches(-0.25)
            p_ref.paragraph_format.space_after = Pt(4)
            p_ref.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = p_ref.add_run(line[2:].strip())
            r.font.size = Pt(9.5)
        elif line.startswith("a. ") or line.startswith("b. ") or line.startswith("c. ") or line.startswith("d. "):
            p_li = doc.add_paragraph()
            p_li.paragraph_format.left_indent = Inches(0.25)
            p_li.paragraph_format.space_after = Pt(3)
            p_li.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = p_li.add_run(line)
            r.font.size = Pt(10.5)
        elif line.startswith("```"):
            # code block skip or render as verbatim
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            p_c = doc.add_paragraph()
            p_c.paragraph_format.left_indent = Inches(0.3)
            p_c.paragraph_format.space_before = Pt(4)
            p_c.paragraph_format.space_after = Pt(6)
            rc = p_c.add_run("\n".join(code_lines))
            rc.font.name = 'Consolas'
            rc.font.size = Pt(8)
        else:
            p_body = doc.add_paragraph()
            p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_body.paragraph_format.space_after = Pt(6)
            p_body.paragraph_format.line_spacing = 1.15
            p_body.paragraph_format.first_line_indent = Inches(0.3)
            # Remove inline math and basic markdown markers for clean reading
            clean_line = line.replace("$$", "").replace("$", "")
            r = p_body.add_run(clean_line)
            r.font.size = Pt(10.5)
            
        i += 1
        
    if table_rows:
        process_table(table_rows)

    doc.save(output_docx_path)
    print(f"File Word SINTA 2 berhasil dibuat: {output_docx_path}")

if __name__ == "__main__":
    create_sinta2_docx(
        paper_md_path="paper.md",
        output_docx_path="tax_compliance_manuscript_sinta2.docx"
    )
