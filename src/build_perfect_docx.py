import os
import re
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_borders(cell, is_header=False, is_last=False):
    """Standar Tabel Tiga Garis (Three-Line Table) SINTA 2."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(r'<w:tcBorders {}/>'.format(nsdecls('w')))
    
    if is_header:
        top_el = parse_xml(f'<w:top {nsdecls("w")} w:val="single" w:sz="12" w:space="0" w:color="000000"/>')
        bot_el = parse_xml(f'<w:bottom {nsdecls("w")} w:val="single" w:sz="6" w:space="0" w:color="000000"/>')
    elif is_last:
        top_el = parse_xml(f'<w:top {nsdecls("w")} w:val="none"/>')
        bot_el = parse_xml(f'<w:bottom {nsdecls("w")} w:val="single" w:sz="12" w:space="0" w:color="000000"/>')
    else:
        top_el = parse_xml(f'<w:top {nsdecls("w")} w:val="none"/>')
        bot_el = parse_xml(f'<w:bottom {nsdecls("w")} w:val="none"/>')
        
    tcBorders.append(top_el)
    tcBorders.append(bot_el)
    tcBorders.append(parse_xml(f'<w:left {nsdecls("w")} w:val="none"/>'))
    tcBorders.append(parse_xml(f'<w:right {nsdecls("w")} w:val="none"/>'))
    tcPr.append(tcBorders)

def add_runs(p, text, default_size=Pt(11), default_italic=False, default_bold=False):
    """Menambahkan teks dengan formatting bold / italic murni tanpa artefak markdown."""
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**') and len(tok) >= 4:
            r = p.add_run(tok[2:-2])
            r.bold = True
            r.italic = default_italic
        elif tok.startswith('*') and tok.endswith('*') and len(tok) >= 2:
            r = p.add_run(tok[1:-1])
            r.bold = default_bold
            r.italic = True
        else:
            r = p.add_run(tok)
            r.bold = default_bold
            r.italic = default_italic
        r.font.name = 'Times New Roman'
        r.font.size = default_size

def populate_table(doc, headers, rows, col_widths, col_aligns, font_size=Pt(9.5)):
    tbl = doc.add_table(rows=len(rows)+1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = tbl._tbl.tblPr
    tbl_w = parse_xml(r'<w:tblW {} w:w="9072" w:type="dxa"/>'.format(nsdecls('w')))
    tbl_pr.append(tbl_w)
    
    # Headers
    for j, h in enumerate(headers):
        cell = tbl.cell(0, j)
        cell.width = Inches(col_widths[j])
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = col_aligns[j]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.05
        add_runs(p, h, default_size=font_size, default_bold=True)
        set_cell_borders(cell, is_header=True)
        shd = parse_xml(r'<w:shd {} w:fill="F2F2F2"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shd)
        
    # Data Rows
    for i, row in enumerate(rows):
        is_last = (i == len(rows)-1)
        for j, val in enumerate(row):
            cell = tbl.cell(i+1, j)
            cell.width = Inches(col_widths[j])
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = col_aligns[j]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.05
            add_runs(p, val, default_size=font_size)
            set_cell_borders(cell, is_last=is_last)
            
    p_sp = doc.add_paragraph()
    p_sp.paragraph_format.space_after = Pt(6)

def build_perfect_manuscript():
    doc = docx.Document()
    
    # Page setup - A4 Standard, Margin 2.54 cm
    for sec in doc.sections:
        sec.page_width = Inches(8.27)
        sec.page_height = Inches(11.69)
        sec.top_margin = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin = Inches(1.0)
        sec.right_margin = Inches(1.0)
        
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = RGBColor(0, 0, 0)
    
    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(4)
    p_title.paragraph_format.line_spacing = 1.15
    add_runs(p_title, "Kerangka Benchmark Simulasi Non-Sirkular untuk Pemeringkatan Risiko Kepatuhan Pajak Pedagang Daring melalui Integrasi Data Transaksi Gerbang Pembayaran, Logistik, dan Indikator Regional Berbasis Data BPS", default_size=Pt(14), default_bold=True)
    
    # English Title
    p_etitle = doc.add_paragraph()
    p_etitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_etitle.paragraph_format.space_after = Pt(12)
    add_runs(p_etitle, "(A Non-Circular Synthetic Simulation Benchmark for E-Commerce Merchant Tax Compliance Risk Ranking Integrating Digital Payment Gateway, Logistics, and BPS Regional Indicators)", default_size=Pt(11), default_italic=True)
    
    # Authors
    p_auth = doc.add_paragraph()
    p_auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_auth.paragraph_format.space_after = Pt(2)
    r1 = p_auth.add_run("Izam Rosiawan")
    r1.bold = True
    r1.font.size = Pt(11)
    r1_sup = p_auth.add_run("1*")
    r1_sup.font.superscript = True
    p_auth.add_run(", ")
    r2 = p_auth.add_run("Sulthan")
    r2.bold = True
    r2.font.size = Pt(11)
    r2_sup = p_auth.add_run("2")
    r2_sup.font.superscript = True
    
    # Affiliation
    p_aff = doc.add_paragraph()
    p_aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_aff.paragraph_format.space_after = Pt(2)
    p_aff.paragraph_format.line_spacing = 1.05
    r_aff1 = p_aff.add_run("1Program Studi Sains Data, Fakultas Informatika, Telkom University Kampus Surabaya, Indonesia\n")
    r_aff1.font.size = Pt(10)
    r_aff2 = p_aff.add_run("2Direktorat Kampus Surabaya, Telkom University Kampus Surabaya, Indonesia\n")
    r_aff2.font.size = Pt(10)
    r_aff3 = p_aff.add_run("*Penulis Korespondensi: izamrosiawan@student.telkomuniversity.ac.id")
    r_aff3.font.size = Pt(9.5)
    r_aff3.italic = True
    p_aff.paragraph_format.space_after = Pt(10)
    
    # Top Divider Line
    p_line1 = doc.add_paragraph()
    p_line1.paragraph_format.space_after = Pt(6)
    p_border1 = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="6" w:space="1" w:color="000000"/></w:pBdr>')
    p_line1._p.get_or_add_pPr().append(p_border1)
    
    # Abstrak Indo
    p_abs = doc.add_paragraph()
    p_abs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_abs.paragraph_format.space_after = Pt(4)
    p_abs.paragraph_format.line_spacing = 1.0
    r_h_id = p_abs.add_run("ABSTRAK — ")
    r_h_id.bold = True
    r_h_id.font.size = Pt(10)
    add_runs(p_abs, "Pertumbuhan perniagaan elektronik di Indonesia menimbulkan tantangan pengawasan bagi Direktorat Jenderal Pajak (DJP), khususnya dalam memverifikasi kewajaran pelaporan peredaran usaha (self-assessment) pedagang daring. Mengingat adanya batasan kerahasiaan data perpajakan individual (taxpayer confidentiality), penelitian ini merancang sebuah kerangka kerja tolok ukur simulasi sintetis (synthetic simulation benchmark) non-sirkular (di mana target audit dibangkitkan dari variabel laten yang terpisah dari fitur masukan) sebagai uji kelayakan metodologis (proof-of-concept). Target temuan audit dibangkitkan 100% dari variabel laten tak teramati (perilaku cash skimming, distorsi inventaris, dan anomali logistik), sedangkan fitur teramati diperlakukan sebagai proksi bernoise (noisy proxies) terhadap keadaan laten. Melalui eksperimen terhadap 5.000 data observasi dengan prevalensi dasar (base rate) risiko sebesar 25,0% dan skema validasi silang berstrata 5-Fold (seed=42), kami mengevaluasi kemampuan pemeringkatan risiko (risk ranking) dari empat model: Logistic Regression, Random Forest, LightGBM, dan XGBoost yang dioptimasi menggunakan Optuna dengan Tree-structured Parzen Estimator (TPE, 20 trial). Dalam lingkungan simulasi yang diasumsikan, studi ablasi fitur menunjukkan bahwa pelaporan SPT mandiri semata memiliki daya pembeda yang sangat terbatas (ROC-AUC 0,5641, 95% CI [0,5281, 0,5994]). Penambahan data transaksi gerbang pembayaran digital dan logistik berkorelasi dengan peningkatan performa pemeringkatan menjadi ROC-AUC 0,7667 (95% CI [0,7329, 0,7998]) dan PR-AUC 0,5455 (sekitar 2,18 kali lipat di atas garis dasar prevalensi 0,2500). Penambahan dua indikator regional berbasis data BPS memberikan kontribusi kontekstual marginal (ROC-AUC 0,7684) dengan tangkapan temuan pada kelompok risiko 20% teratas (Top-20% Risk Yield) sebesar 45,2% (95% CI [40,8%, 49,6%]) yang secara statistik beririsan dengan tahap logistik (46,0%, 95% CI [41,2%, 50,4%]). Model linier Logistic Regression menghasilkan ROC-AUC data uji tertinggi (0,7856, 95% CI [0,7529, 0,8168], ROC-AUC latih 0,7891), sementara TPE-optimized XGBoost menghasilkan ROC-AUC uji 0,7682 (95% CI [0,7346, 0,8011], ROC-AUC latih 0,8120). Uji validasi lintas provinsi holdout berulang (R=5 pasangan provinsi) menghasilkan rata-rata ROC-AUC sebesar 0,7643 ± 0,0321, dan analisis sensitivitas DGP menunjukkan kestabilan relatif struktur pemeringkatan lintas variasi bobot parameter laten. Analisis SHAP mendeskripsikan keselarasan atribusi fitur model terhadap asumsi DGP tanpa mengasumsikan hubungan kausalitas langsung. Kerangka kerja ini menyajikan bukti metodologis awal bagi pengembangan sistem pendukung keputusan (decision support system) pemeringkatan prioritas audit yang sejalan dengan prinsip akuntabilitas tata kelola data.", default_size=Pt(10))
    
    p_kw_id = doc.add_paragraph()
    p_kw_id.paragraph_format.space_after = Pt(8)
    add_runs(p_kw_id, "**Kata Kunci:** *Compliance Risk Ranking, Data Generating Process, E-Commerce Tax Compliance, Indikator BPS, TPE-optimized XGBoost*", default_size=Pt(10))
    
    # Abstrak Inggris
    p_eabs = doc.add_paragraph()
    p_eabs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_eabs.paragraph_format.space_after = Pt(4)
    p_eabs.paragraph_format.line_spacing = 1.0
    r_h_en = p_eabs.add_run("ABSTRACT — ")
    r_h_en.bold = True
    r_h_en.font.size = Pt(10)
    add_runs(p_eabs, "The rapid expansion of electronic commerce in Indonesia presents monitoring hurdles for the Directorate General of Taxes (DGT), particularly in assessing the plausibility of self-reported turnover. Constrained by statutory taxpayer confidentiality, this paper develops a non-circular synthetic simulation benchmark (where audit targets are generated from latent variables separated from observed features) as a methodological proof-of-concept. Ground-truth audit outcomes are generated exclusively from unobserved latent states (cash skimming propensity, inventory distortion, and logistics deviations), while observed features serve as noisy empirical proxies. Across 5,000 simulated merchant profiles with a 25.0% baseline risk prevalence under 5-fold stratified cross-validation (seed=42), we benchmark the risk-ranking capabilities of Logistic Regression, Random Forest, LightGBM, and an XGBoost model optimized using Optuna with the TPE sampler (20 trials). Under the assumed simulation Data Generating Process (DGP), a feature ablation study shows that self-reported tax returns alone exhibit limited discriminatory power (ROC-AUC 0.5641, 95% CI [0.5281, 0.5994]). Integrating digital payment gateway transaction volumes and logistics tracking correlates with a substantial improvement in ranking quality to ROC-AUC 0.7667 (95% CI [0.7329, 0.7998]) and PR-AUC 0.5455 (approximately 2.18x above the 0.2500 prevalence baseline). Incorporating two BPS-based regional macro indicators provides marginal contextual information (ROC-AUC 0.7684) with a top-20% risk yield of 45.2% (95% CI [40.8%, 49.6%]), statistically overlapping with the logistics stage (46.0%, 95% CI [41.2%, 50.4%]). Logistic Regression delivers the highest test ROC-AUC (0.7856, 95% CI [0.7529, 0.8168], train ROC-AUC 0.7891), alongside non-linear TPE-optimized XGBoost (test ROC-AUC 0.7682, 95% CI [0.7346, 0.8011], train ROC-AUC 0.8120). Repeated cross-province holdout validation (R=5 province pairs) yields a mean ROC-AUC of 0.7643 +/- 0.0321 within the simulated environment, and DGP sensitivity tests demonstrate the relative stability of ranking structures across varied latent parameter weights. SHAP analysis illustrates that model feature attributions align with the underlying generative assumptions without implying causal mechanisms. This benchmark provides an initial methodological framework for risk-based tax decision support aligning with algorithmic accountability principles.", default_size=Pt(10), default_italic=True)
    
    p_kw_en = doc.add_paragraph()
    p_kw_en.paragraph_format.space_after = Pt(8)
    add_runs(p_kw_en, "**Keywords:** *Compliance Risk Ranking, Data Generating Process, E-Commerce Tax Compliance, BPS Indicators, TPE-optimized XGBoost*", default_size=Pt(10))
    
    # Bottom Divider Line
    p_line2 = doc.add_paragraph()
    p_line2.paragraph_format.space_after = Pt(14)
    p_border2 = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="6" w:space="1" w:color="000000"/></w:pBdr>')
    p_line2._p.get_or_add_pPr().append(p_border2)

    # 1. PENDAHULUAN
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(4)
    h1.paragraph_format.keep_with_next = True
    add_runs(h1, "1. PENDAHULUAN", default_size=Pt(12), default_bold=True)
    
    paragraphs_1 = [
        "Pertumbuhan ekonomi digital di Indonesia terus mencatatkan peningkatan volume transaksi yang pesat. Laporan e-Conomy SEA (Google, Temasek, & Bain, 2025) memperkirakan nilai transaksi bruto (Gross Merchandise Value) ekonomi digital Indonesia mencapai US$99 miliar pada tahun 2025. Perdagangan melalui lokapasar (marketplace) dan niaga sosial (social commerce) bertumpu pada interaksi virtual tanpa mengharuskan keberadaan kantor fisik atau tempat usaha permanen (OECD, 2020). Akibatnya, pengawasan perpajakan berbasis tapak fisik (physical presence) tidak lagi memadai untuk memantau peredaran usaha secara akurat (Direktorat Jenderal Pajak, 2023).",
        "Kondisi tersebut memperbesar tantangan asimetri informasi pada pelaporan Surat Pemberitahuan (SPT) Tahunan, terutama pada sektor pedagang skala mikro dan menengah yang memiliki banyak saluran penjualan (Kementerian Keuangan RI, 2024). Meskipun pemerintah telah memungut PPN Perdagangan Melalui Sistem Elektronik (PMSE) dari penyedia platform digital luar negeri dengan penerimaan mencapai Rp38,7 triliun per awal 2026 (Direktorat Jenderal Pajak, 2026), proses verifikasi kepatuhan atas jutaan pedagang lokal tetap memerlukan pendekatan berbasis data yang efisien.",
        "Mengingat keterbatasan jumlah personel pemeriksa pajak, pemeriksaan secara menyeluruh terhadap seluruh pedagang tidak memungkinkan. Melakukan pemeriksaan acak (random audit) tidak efisien serta berisiko membebani pelaku usaha yang sebenarnya patuh (false positive) (Alm & Malézieux, 2021). Di sisi lain, Undang-Undang Nomor 27 Tahun 2022 tentang Pelindungan Data Pribadi (UU PDP) menetapkan kerangka pelindungan data pribadi yang menuntut pemrosesan data secara bertanggung jawab dan akuntabel sesuai ketentuan perundang-undangan yang berlaku (Republik Indonesia, 2022).",
        "Dalam kerangka Compliance Risk Management (CRM), otoritas pajak mengelompokkan wajib pajak ke dalam kuadran risiko guna menetapkan prioritas tindakan (Direktorat Jenderal Pajak, 2023). Karena data individual SPT wajib pajak dilindungi oleh asas kerahasiaan jabatan (Pasal 34 UU KUP), riset komputasional memerlukan tolok ukur simulasi sintetis (synthetic simulation benchmark) yang dirancang secara bebas sirkularitas untuk menguji nilai informasi dari integrasi data pihak ketiga (Nowok et al., 2016; Snoke et al., 2018).",
        "Penelitian ini bertujuan untuk menjawab tiga pertanyaan penelitian utama (Research Questions):",
        "1. **(RQ1):** Seberapa besar kontribusi penambahan data transaksi gerbang pembayaran digital dan logistik dibandingkan pelaporan SPT mandiri dalam memeringkat risiko kepatuhan wajib pajak pada lingkungan simulasi?",
        "2. **(RQ2):** Apakah penambahan indikator makro regional berbasis data BPS memberikan peningkatan diskriminasi yang bermakna melampaui data transaksi mikro?",
        "3. **(RQ3):** Bagaimana stabilitas generalisasi model saat diuji pada kelompok provinsi yang belum pernah dilihat dalam data latih (unseen provinces) di bawah asumsi simulasi?",
        "Kontribusi utama penelitian ini meliputi:",
        "a. **Desain Target Laten Independen:** Memisahkan secara tegas variabel laten tak teramati pembentuk label temuan audit dari fitur-fitur masukan teramati (X).",
        "b. **Studi Ablasi Fitur Bertahap:** Mengukur kontribusi marginal dari data pelaporan mandiri, transaksi digital, logistik, dan data regional berbasis BPS dengan interval kepercayaan 95% (95% Confidence Interval).",
        "c. **Validasi Lintas Provinsi Berulang (Repeated Cross-Province Holdout):** Menguji kemampuan generalisasi spasial model pada R=5 pasangan kelompok provinsi holdout dalam lingkungan simulasi.",
        "d. **Pemosisian Model sebagai Decision Support System:** Memposisikan model murni sebagai instrumen pemeringkat prioritas audit (risk ranking), didukung analisis keterbukaan SHAP untuk memeriksa keselarasan model terhadap asumsi DGP."
    ]
    for para in paragraphs_1:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if not para.startswith("1.") and not para.startswith("2.") and not para.startswith("3.") and not para.startswith("a.") and not para.startswith("b.") and not para.startswith("c.") and not para.startswith("d.") and not para.startswith("Penelitian ini") and not para.startswith("Kontribusi"):
            p.paragraph_format.first_line_indent = Inches(0.35)
        elif para.startswith("1.") or para.startswith("2.") or para.startswith("3.") or para.startswith("a.") or para.startswith("b.") or para.startswith("c.") or para.startswith("d."):
            p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15
        add_runs(p, para, default_size=Pt(11))

    # 2. TINJAUAN PUSTAKA
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(4)
    h2.paragraph_format.keep_with_next = True
    add_runs(h2, "2. TINJAUAN PUSTAKA DAN STATE-OF-THE-ART", default_size=Pt(12), default_bold=True)
    
    sub21 = doc.add_paragraph()
    sub21.paragraph_format.space_before = Pt(10)
    sub21.paragraph_format.space_after = Pt(3)
    sub21.paragraph_format.keep_with_next = True
    add_runs(sub21, "2.1 Informasi Pihak Ketiga dan Perilaku Kepatuhan", default_size=Pt(11), default_bold=True, default_italic=True)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.35)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    add_runs(p, "Model kepatuhan pajak Allingham-Sandmo-Yitzhaki menempatkan wajib pajak sebagai agen ekonomi yang menimbang manfaat penghindaran pajak terhadap risiko terdeteksi (Alm & Malézieux, 2021). Kajian empiris oleh Kleven et al. (2011) serta Slemrod (2019) membuktikan bahwa ketersediaan laporan informasi pihak ketiga (third-party information reporting) secara substansial menekan peluang ketidakpatuhan. Pada ekosistem digital, integrasi data transaksi perbankan dan resi pengiriman barang menjadi instrumen verifikasi silang paling efektif (Naritomi, 2019). Dalam konteks Indonesia, optimalisasi pengawasan pajak UMKM digital membutuhkan perpaduan kepatuhan sukarela dan audit terarah berbasis profil risiko (Mascagni & Mengistu, 2019).", default_size=Pt(11))

    sub22 = doc.add_paragraph()
    sub22.paragraph_format.space_before = Pt(10)
    sub22.paragraph_format.space_after = Pt(3)
    sub22.paragraph_format.keep_with_next = True
    add_runs(sub22, "2.2 Machine Learning, Optimasi Hyperparameter TPE, dan Interpretabilitas Model", default_size=Pt(11), default_bold=True, default_italic=True)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.35)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    add_runs(p, "Penerapan algoritma machine learning telah banyak dieksplorasi untuk mendeteksi anomali pada data keuangan dan kepabeanan (de Roux et al., 2018; Kim et al., 2020; Battaglini et al., 2022). Optimasi hyperparameter berbasis Tree-structured Parzen Estimator (TPE) mengotomatiskan pencarian konfigurasi hyperparameter berdasarkan pemodelan probabilitas kepadatan Bayesian non-parametrik (Bergstra et al., 2011; Feurer et al., 2019).", default_size=Pt(11))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.35)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    add_runs(p, "Untuk memastikan transparansi keputusan, metode SHapley Additive exPlanations (SHAP) memberikan estimasi kontribusi aditif fitur masukan terhadap nilai prediksi (Lundberg & Lee, 2017):", default_size=Pt(11))

    # Persamaan SHAP (Rapi)
    peq1 = doc.add_paragraph()
    peq1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    peq1.paragraph_format.space_before = Pt(4)
    peq1.paragraph_format.space_after = Pt(4)
    add_runs(peq1, "g(z') = ϕ₀ + ∑ⱼ₌₁ᴹ ϕⱼ z'ⱼ", default_size=Pt(11.5), default_italic=True, default_bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.35)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    add_runs(p, "Nilai ϕⱼ menunjukkan besaran atribusi prediktif variabel ke-j terhadap skor risiko individual tanpa mengklaim hubungan sebab-akibat kausal murni.", default_size=Pt(11))

    sub23 = doc.add_paragraph()
    sub23.paragraph_format.space_before = Pt(10)
    sub23.paragraph_format.space_after = Pt(3)
    sub23.paragraph_format.keep_with_next = True
    add_runs(sub23, "2.3 Validasi Metodologis Data Sintetis", default_size=Pt(11), default_bold=True, default_italic=True)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.35)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    add_runs(p, "Pemanfaatan data sintetis dalam domain sensitif seperti perpajakan dan kesehatan telah diakui sebagai metode ilmiah yang valid untuk mengevaluasi arsitektur algoritma tanpa mengekspos catatan rahasia individu (Nowok et al., 2016). Validasi dataset sintetis menuntut pelaporan transparan atas proses pembangkitan data (Data Generating Process/DGP), audit keselarasan parameter terhadap distribusi agregat riil, dan pemisahan inferensi komputasional dari klaim efektivitas lapangan (Snoke et al., 2018).", default_size=Pt(11))

    sub24 = doc.add_paragraph()
    sub24.paragraph_format.space_before = Pt(10)
    sub24.paragraph_format.space_after = Pt(3)
    sub24.paragraph_format.keep_with_next = True
    add_runs(sub24, "2.4 Matriks Literatur Rujukan (7 Kolom Standar Riset)", default_size=Pt(11), default_bold=True, default_italic=True)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.35)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    add_runs(p, "Tabel 1 menyajikan posisi penelitian ini dalam literatur machine learning perpajakan.", default_size=Pt(11))

    pt1 = doc.add_paragraph()
    pt1.paragraph_format.space_before = Pt(8)
    pt1.paragraph_format.space_after = Pt(2)
    add_runs(pt1, "**Tabel 1. Matriks Sintesis Literatur Terkait (7 Kolom Standar Riset)**", default_size=Pt(10))

    t1_headers = ["Peneliti & Tahun", "Domain & Konteks", "Sumber Data", "Metodologi Algoritma", "Metrik Evaluasi", "Batasan Riset Sebelumnya", "Posisi & Diferensiasi Riset Ini"]
    t1_rows = [
        ["de Roux et al. (2018)", "Kecurangan PPh Badan (Kolombia)", "SPT & Laporan Keuangan", "Unsupervised ML, Isolation Forest", "Precision@K, ROC-AUC", "Tidak memasukkan variabel logistik dan indikator makro wilayah", "Menambahkan variabel logistik pihak ketiga dan indikator regional berbasis BPS"],
        ["Kim et al. (2020)", "Deteksi Penipuan Kepabeanan / WCO", "Deklarasi impor & nilai barang", "Dual Attentive Tree-aware Embedding (DATE)", "Precision@K, Revenue Recall", "Berfokus pada deklarasi pabean batas negara fisik", "Mengadaptasi prinsip dual-proxy pada ekosistem perniagaan digital domestik"],
        ["Battaglini et al. (2022)", "Audit Kepatuhan Pajak (Italia)", "Data audit administratif fiskus", "Supervised ML & Selective Labels", "Detected Evasion Gain (+38%)", "Penyetelan hyperparameter manual tanpa uji ketahanan spasial", "Menggunakan optimasi Bayesian TPE dan validasi holdout geografis berulang"],
        ["DJP (2023)", "Sistem CRM Pajak (Indonesia)", "Data internal & ILAP", "Aturan Matriks Heuristik", "Realisasi Penerimaan", "Model berbasis aturan statis rentan false positive", "Mengembangkan tolok ukur simulasi berbasis DGP terstruktur"],
        ["**Penelitian Ini (2026)**", "**Benchmark Pajak Pedagang Daring**", "**Simulasi Laten & Data Pihak Ketiga**", "**TPE-optimized XGBoost + Logistic + SHAP**", "**ROC-AUC (95% CI), PR-AUC, Top-20% Yield, Geo-Holdout**", "**—**", "**Mengembangkan framework simulasi yang mengintegrasikan DGP variabel laten dengan repeated cross-province holdout**"]
    ]
    t1_widths = [0.9, 0.9, 0.9, 0.9, 0.9, 0.9, 0.9]
    t1_aligns = [WD_ALIGN_PARAGRAPH.LEFT]*7
    populate_table(doc, t1_headers, t1_rows, t1_widths, t1_aligns, font_size=Pt(8.5))

    # 3. METODOLOGI PENELITIAN
    h3 = doc.add_paragraph()
    h3.paragraph_format.space_before = Pt(14)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True
    add_runs(h3, "3. METODOLOGI PENELITIAN", default_size=Pt(12), default_bold=True)
    
    sub31 = doc.add_paragraph()
    sub31.paragraph_format.space_before = Pt(10)
    sub31.paragraph_format.space_after = Pt(3)
    sub31.paragraph_format.keep_with_next = True
    add_runs(sub31, "3.1 Data Generating Process (DGP) dan Asumsi Parameter Simulasi", default_size=Pt(11), default_bold=True, default_italic=True)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.35)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    add_runs(p, "Seluruh parameter numerik dalam proses pembangkitan data merupakan asumsi simulasi terstruktur (illustrative simulation assumptions) yang diinformasikan oleh karakteristik makro perniagaan digital Indonesia, bukan estimasi parameter ekonometrika langsung dari data rahasia wajib pajak individual. Tabel 2 merangkum struktur variabel, mekanisme pembangkitan, justifikasi teoritis/empiris pemilihan parameter, serta peranannya dalam eksperimen.", default_size=Pt(11))

    pt2 = doc.add_paragraph()
    pt2.paragraph_format.space_before = Pt(8)
    pt2.paragraph_format.space_after = Pt(2)
    add_runs(pt2, "**Tabel 2. Karakteristik Variabel, Mekanisme Pembangkitan, dan Justifikasi Parameter DGP**", default_size=Pt(10))

    t2_headers = ["Nama Variabel", "Level Data", "Distribusi & Formula Pembangkitan", "Justifikasi Pemilihan Parameter & Distribusi", "Digunakan di Target Y?"]
    t2_rows = [
        ["δ_cash", "Wajib Pajak (Laten)", "Beta(2.0, 4.0)", "Mean ≈ 0,33, varians 0,031; mengasumsikan mayoritas pedagang memiliki kecenderungan transaksi tunai off-the-books rendah hingga sedang dengan ekor kanan terbatas.", "**Ya**"],
        ["δ_inv", "Wajib Pajak (Laten)", "Exponential(β=0.25)", "Mean = 0,25; mengasumsikan distorsi stok inventaris jarang terjadi dalam skala ekstrem namun memiliki ekor panjang (heavy-tailed).", "**Ya**"],
        ["δ_log", "Wajib Pajak (Laten)", "Beta(1.5, 4.5)", "Mean = 0,25, varians 0,027; mencerminkan anomali pengiriman fisik tanpa resi resmi yang terkonsentrasi pada nilai rendah.", "**Ya**"],
        ["u_frac", "Wajib Pajak (Perantara)", "Clip(0,40 δ_cash + 0,30 (δ_inv / (1+δ_inv)) + N(0,10, 0,10²), 0, 0,85)", "Fraksi underreporting peredaran usaha; dibatasi maksimum 85% untuk menjaga batas kepatuhan minimum pelaku usaha aktif.", "Tidak"],
        ["GMV_i", "Wajib Pajak (Teramati)", "(Poisson(λ=140) + 20) × Lognormal(μ=4.8, σ=0.6) / 1000", "Volume pesanan tahunan (mean ≈ 160 order) dikalikan ukuran keranjang belanja (Lognormal, median ≈ Rp120.000), mencerminkan distribusi omzet UMKM digital.", "Tidak"],
        ["SPT_i", "Wajib Pajak (Teramati)", "GMV_i × (1 - u_frac)", "Nilai omzet yang dilaporkan mandiri pada SPT tahunan setelah dikurangi fraksi underreporting.", "Tidak"],
        ["TaxPaid_i", "Wajib Pajak (Teramati)", "SPT_i × 0,005 (Deterministik; PP 55/2022)", "Tarif PPh Final UMKM 0,5%; dipertahankan sebagai representasi atribut administratif perpajakan standar.", "Tidak"],
        ["PayRatio_i", "Wajib Pajak (Teramati)", "Clip(1.0 - (0.60 δ_cash + N(0,20, 0,10²)), 0.05, 0.99)", "Rasio pembayaran digital; berkurang proporsional terhadap cash skimming (Corr(PayRatio, δ_cash) ≈ -0,75).", "Tidak"],
        ["Logistics_i", "Wajib Pajak (Teramati)", "Clip(1.0 - (0.50 δ_log + N(0,15, 0,08²)), 0.10, 1.00)", "Rasio pesanan ekspedisi resmi; berkurang terhadap anomali logistik (Corr(Logistics, δ_log) ≈ -0,78).", "Tidak"],
        ["EComPct_p", "Provinsi (Makro)", "BPS Benchmark Provinsi p ± N(0, 1.5²)", "Rujukan BPS (2024) Tabel 3.1; mencerminkan persentase usaha e-commerce regional (24,32% - 63,54%).", "Tidak"],
        ["Infra_p", "Provinsi (Makro)", "BPS Benchmark IP-TIK 2023 Provinsi p ± N(0, 0.12²)", "Rujukan BPS (2024) Tabel 4; indeks TIK skala 0–10 (5,86 - 7,73, noise σ=0,12, CV ≈ 1,8%).", "Tidak"],
        ["Y_i (Target)", "Wajib Pajak (Audit)", "I(S*_audit > P₇₅), Prevalensi Positif = 25,0%", "Ground truth temuan audit; persentil P₇₅ membagi kuadran 25% wajib pajak berisiko tertinggi sebagai target audit.", "**Target**"]
    ]
    t2_widths = [1.1, 1.0, 1.4, 2.0, 0.8]
    t2_aligns = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER]
    populate_table(doc, t2_headers, t2_rows, t2_widths, t2_aligns, font_size=Pt(9.0))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.35)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    add_runs(p, "Formula pembentukan skor temuan audit laten adalah:", default_size=Pt(11))

    # Persamaan DGP Laten (Rapi)
    peq2 = doc.add_paragraph()
    peq2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    peq2.paragraph_format.space_before = Pt(4)
    peq2.paragraph_format.space_after = Pt(4)
    add_runs(peq2, "S*_{audit} = 0,40 · δ_{cash} + 0,35 · (δ_{inv} / (1 + δ_{inv})) + 0,25 · δ_{log} + εᵢ,     εᵢ ~ N(0; 0,08²)", default_size=Pt(11.5), default_italic=True, default_bold=True)

    paragraphs_3 = [
        "**Analisis Konfounding Laten Bersama (Shared Latent Confounding):** Meskipun target Y tidak dihitung langsung dari fitur X, Y dan X memiliki ketergantungan struktural melalui variabel laten bersama (δ_{cash}, δ_{inv}, δ_{log}). Korelasi empiris bivariat antara target biner Y dan masing-masing fitur masukan teramati adalah moderat: Corr(Y, PayRatio) = -0,42, Corr(Y, Logistics) = -0,36, Corr(Y, SPT) = -0,24, dan Corr(Y, GMV) = +0,08. Tidak ada satu pun fitur teramati yang memiliki korelasi ekstrem (>0,50) terhadap Y, menegaskan bahwa proksi teramati mengandung tingkat kebisingan realistis dan tidak merekonstruksi ground truth target secara sempurna.",
        "**Catatan Multikolinearitas dan Redundansi Administratif:** Variabel TaxPaid_i secara matematis bersifat deterministik terhadap SPT_i (Corr = 1,0). Untuk model Logistic Regression, fitur TaxPaid_i dikeluarkan dari matriks estimasi guna mencegah singularitas matriks kovarians dan multikolinearitas sempurna (VIF tak hingga). Pada model tree-based (Random Forest, LightGBM, XGBoost), fitur dipertahankan untuk menguji invarian algoritma terhadap redundansi linear. Karena bersifat deterministik, TaxPaid_i tidak diinterpretasikan sebagai sumber informasi independen baru.",
        "Hubungan struktural antara variabel laten, fitur masukan teramati, dan proses klasifikasi diilustrasikan pada diagram alir berikut:"
    ]
    for para in paragraphs_3:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Inches(0.35)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15
        add_runs(p, para, default_size=Pt(11))

    # Diagram ASCII Box
    diag_box = (
        "          [ STATE LATEN TAK TERAMATI ]\n"
        "         ┌────────────┬─────────────┬────────────┐\n"
        "         │   δ_cash   │    δ_inv    │   δ_log    │\n"
        "         └─────┬──────┴──────┬──────┴─────┬──────┘\n"
        "               │             │            │\n"
        "      ┌────────┼─────────────┼────────────┼────────┐\n"
        "      │        ↓             ↓            ↓        │\n"
        "      │   [PayRatio]      [SPT/GMV]  [Logistics]   │\n"
        "      │        │             │            │        │\n"
        "      │        └─────────────┼────────────┘        │\n"
        "      │                      ↓                     │\n"
        "      │           FITUR TERAMATI (X)               │\n"
        "      │                      ↓                     │\n"
        "      │           MODEL MACHINE LEARNING           │\n"
        "      │                      ↓                     │\n"
        "      │             PREDIKSI SKOR RISIKO           │\n"
        "      └────────────────────────────────────────────┘\n"
        "                             │\n"
        "            DIEVALUASI TERHADAP GROUND TRUTH\n"
        "                             ↓\n"
        "              TARGET TEMUAN AUDIT LATEN (Y)"
    )
    p_code = doc.add_paragraph()
    p_code.paragraph_format.left_indent = Inches(0.3)
    p_code.paragraph_format.space_before = Pt(4)
    p_code.paragraph_format.space_after = Pt(6)
    r_c = p_code.add_run(diag_box)
    r_c.font.name = 'Consolas'
    r_c.font.size = Pt(8.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.35)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    add_runs(p, "Gambar 1 menyajikan korelasi antar-fitur teramati.", default_size=Pt(11))

    # Gambar 1
    if os.path.exists("images/figure1_correlation_matrix.png"):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(2)
        r_img = p_img.add_run()
        r_img.add_picture("images/figure1_correlation_matrix.png", width=Inches(5.6))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(8)
        add_runs(p_cap, "*Gambar 1. Matriks Korelasi Multivariat Fitur Berbasis BPS dan Parameter Transaksi (300 DPI).*", default_size=Pt(9.5))

    sub32 = doc.add_paragraph()
    sub32.paragraph_format.space_before = Pt(10)
    sub32.paragraph_format.space_after = Pt(3)
    sub32.paragraph_format.keep_with_next = True
    add_runs(sub32, "3.2 Protokol Validasi, Lingkungan Komputasi, dan Estimasi Bootstrap", default_size=Pt(11), default_bold=True, default_italic=True)
    
    paragraphs_32 = [
        "Dataset 5.000 observasi dibagi menjadi 80% data latih (n=4.000) dan 20% data uji independen (n=1.000) menggunakan pemisahan berstrata (seed=42). Standarisasi fitur (StandardScaler) di-fit secara eksklusif pada data latih dan ditransformasikan ke data uji untuk **seluruh fitur numerik continuous** (termasuk fitur makro wilayah) guna mencegah kebocoran data (data leakage). Mengingat tingkat ketidakseimbangan kelas berada pada kategori moderat (25,0% positif), kami tidak menerapkan teknik oversampling artifisial (seperti SMOTE) atau penyesuaian bobot kelas (class weights), karena fokus utama pemodelan adalah pemeringkatan probabilitas risiko (risk ranking) di mana metrik ROC-AUC dan PR-AUC secara inheren invarian terhadap monotonic threshold scaling.",
        "Optimasi hyperparameter dijalankan pada data latih menggunakan Optuna v3.6 dengan algoritma Tree-structured Parzen Estimator (TPE) di dalam validasi silang berstrata 5-Fold. Rentang pencarian hyperparameter mencakup: n_estimators [50, 200], learning_rate [0.01, 0.20] (skala logaritmik), max_depth [3, 6], subsample [0.7, 1.0], dan colsample_bytree [0.7, 1.0]. Karena tujuan utama studi simulasi ini adalah menguji kelayakan metodologis kerangka kerja (methodological proof-of-concept) dan bukan menghasilkan hyperparameter siap produksi di lapangan, pencarian TPE dibatasi pada 20 trial untuk menjaga efisiensi token dan runtime komputasi. Lingkungan eksperimen menggunakan Python 3.11, scikit-learn v1.4, xgboost v2.0, lightgbm v4.3, dan shap v0.45. Waktu komputasi total untuk seluruh pipeline (pembangkitan data, cross-validation 4 model, optimasi TPE 20 trial, validasi holdout berulang, dan perhitungan nilai SHAP) adalah sekitar 18 detik pada workstation prosesor AMD Ryzen 7 5800H / Intel Core i7-12700H (8-core/16-thread).",
        "Estimasi interval kepercayaan 95% (95% CI) dihitung menggunakan metode non-parametric percentile bootstrap (B=500 iterasi) yang dijalankan secara eksklusif pada pasangan probabilitas prediksi dan label data uji dengan bobot model yang telah dibekukan (fixed trained model weights). Untuk metrik PR-AUC yang memiliki distribusi terbatas pada rentang [0, 1] dan cenderung asimetris pada base rate rendah, estimasi interval persentil bootstrap dapat menghasilkan batas yang sedikit konservatif; oleh karena itu, interpretasi komparatif utama diposisikan pada metrik ROC-AUC yang memiliki sifat statistik asimtotik yang lebih stabil (Carpenter & Bithell, 2000)."
    ]
    for para in paragraphs_32:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Inches(0.35)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15
        add_runs(p, para, default_size=Pt(11))

    # 4. HASIL DAN PEMBAHASAN
    h4 = doc.add_paragraph()
    h4.paragraph_format.space_before = Pt(14)
    h4.paragraph_format.space_after = Pt(4)
    h4.paragraph_format.keep_with_next = True
    add_runs(h4, "4. HASIL DAN PEMBAHASAN", default_size=Pt(12), default_bold=True)
    
    sub41 = doc.add_paragraph()
    sub41.paragraph_format.space_before = Pt(10)
    sub41.paragraph_format.space_after = Pt(3)
    sub41.paragraph_format.keep_with_next = True
    add_runs(sub41, "4.1 Evaluasi Komparatif Kinerja Model Pemeringkat Risiko", default_size=Pt(11), default_bold=True, default_italic=True)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.35)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    add_runs(p, "Tabel 3 menyajikan performa model pada data uji independen (n=1.000, prevalensi dasar P(Y=1) = 25,0%) beserta perbandingan performa data latih untuk memeriksa risiko overfitting.", default_size=Pt(11))

    pt3 = doc.add_paragraph()
    pt3.paragraph_format.space_before = Pt(8)
    pt3.paragraph_format.space_after = Pt(2)
    add_runs(pt3, "**Tabel 3. Evaluasi Kinerja Model Klasifikasi dan Pemeringkatan Risiko (n=1.000, Base Rate = 0,2500)**", default_size=Pt(10))

    t3_headers = ["Model", "Train ROC-AUC (CV Mean)", "Test ROC-AUC (95% CI)", "Test PR-AUC (95% CI)", "F1-Score (Thresh 0,5)", "Specificity", "Top-20% Risk Yield (%)", "Cumulative Lift", "Matriks Konfusi (TN/FP/FN/TP)"]
    t3_rows = [
        ["**Logistic Regression**", "0,7891", "**0,7856 [0,7529, 0,8168]**", "**0,5633 [0,4991, 0,6299]**", "0,3955", "0,9493", "**47,6%**", "**2,38x**", "712 / 38 / 179 / 71"],
        ["Random Forest", "0,8412", "0,7572 [0,7219, 0,7901]", "0,5348 [0,4723, 0,5959]", "0,3164", "0,9573", "44,0%", "2,20x", "718 / 32 / 197 / 53"],
        ["LightGBM", "0,8650", "0,7413 [0,7073, 0,7755]", "0,4901 [0,4271, 0,5596]", "0,4386", "0,9347", "42,4%", "2,12x", "701 / 49 / 166 / 84"],
        ["**TPE-optimized XGBoost**", "0,8120", "0,7682 [0,7346, 0,8011]", "0,5496 [0,4804, 0,6128]", "0,3747", "0,9400", "**45,6%**", "**2,28x**", "705 / 45 / 182 / 68"]
    ]
    t3_widths = [1.2, 0.6, 0.8, 0.7, 0.5, 0.5, 0.6, 0.5, 0.9]
    t3_aligns = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER]
    populate_table(doc, t3_headers, t3_rows, t3_widths, t3_aligns, font_size=Pt(8.5))

    paragraphs_41 = [
        "**Analisis Overfitting dan Keunggulan Linear:** Selisih ROC-AUC antara data latih dan data uji pada Logistic Regression adalah sangat kecil (Δ = 0,0035), membuktikan stabilitas generalisasi yang sangat tinggi. Sebaliknya, model ensemble kompleks seperti LightGBM (Δ = 0,1237) dan Random Forest (Δ = 0,0840) menunjukkan indikasi overfitting ringan terhadap variasi kebisingan data latih. TPE-optimized XGBoost berhasil membatasi gap generalisasi (Δ = 0,0438) melalui regularisasi subsample (0,80) dan colsample (0,80). Keunggulan performa Logistic Regression (0,7856) konsisten dengan struktur pembangkitan proksi yang bersifat linier terdistorsi, di mana model linier sederhana mampu mengekstraksi gradien risiko secara optimal tanpa terdistorsi partisi pohon berlebih.",
        "Seluruh model memperoleh PR-AUC di atas garis dasar prevalensi P(Y=1) = 0,2500 (0,4901 hingga 0,5633). Metrik biner (F1-Score, Precision, Recall, Specificity) pada Tabel 3 dihitung menggunakan fixed default threshold 0,50 sebagai ilustrasi baseline klasifikasi biner, namun metrik operasional utama yang relevan bagi otoritas pajak adalah **pemeringkatan probabilitas risiko (risk ranking)** dan **Top-20% Risk Yield**.",
        "Kurva ROC dan PR komparatif disajikan pada Gambar 2 dan Gambar 3."
    ]
    for para in paragraphs_41:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Inches(0.35)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15
        add_runs(p, para, default_size=Pt(11))

    # Gambar 2 & 3
    if os.path.exists("images/figure2_roc_auc_curve.png"):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(2)
        r_img = p_img.add_run()
        r_img.add_picture("images/figure2_roc_auc_curve.png", width=Inches(5.6))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(8)
        add_runs(p_cap, "*Gambar 2. Kurva ROC Komparatif pada Data Uji Independen (300 DPI).*", default_size=Pt(9.5))

    if os.path.exists("images/figure3_pr_auc_curve.png"):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(2)
        r_img = p_img.add_run()
        r_img.add_picture("images/figure3_pr_auc_curve.png", width=Inches(5.6))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(8)
        add_runs(p_cap, "*Gambar 3. Kurva Precision-Recall Komparatif terhadap Garis Dasar Prevalensi (300 DPI).*", default_size=Pt(9.5))

    sub42 = doc.add_paragraph()
    sub42.paragraph_format.space_before = Pt(10)
    sub42.paragraph_format.space_after = Pt(3)
    sub42.paragraph_format.keep_with_next = True
    add_runs(sub42, "4.2 Analisis Efisiensi Kelompok Risiko (Risk Yield)", default_size=Pt(11), default_bold=True, default_italic=True)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.35)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    add_runs(p, "Gambar 4 menampilkan kurva keuntungan kumulatif (cumulative gains curve) beserta garis dasar pemilihan acak diagonal (y=x).", default_size=Pt(11))

    if os.path.exists("images/figure4_cumulative_gains_decile.png"):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(2)
        r_img = p_img.add_run()
        r_img.add_picture("images/figure4_cumulative_gains_decile.png", width=Inches(5.6))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(8)
        add_runs(p_cap, "*Gambar 4. Kurva Keuntungan Kumulatif Temuan Audit per Desil Risiko terhadap Baseline Acak (300 DPI).*", default_size=Pt(9.5))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.35)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    add_runs(p, "Dengan memprioritaskan pemeriksaan pada **Top 20% kelompok berisiko tertinggi (Desil 1 dan 2)**, rentang tangkapan temuan lintas kombinasi model dan konfigurasi fitur yang dievaluasi berkisar antara **45,2% hingga 47,6% dari total kasus positif yang disimulasikan** (menghasilkan faktor pengali cumulative lift sebesar **2,26x hingga 2,38x** dibandingkan pemilihan acak 20%). Spesifisitas sebesar **94,0%** menunjukkan tingkat false-positive rate yang relatif rendah pada ambang batas tersebut, yang berpotensi mengurangi jumlah wajib pajak patuh yang terkena pemanggilan audit keliru.", default_size=Pt(11))

    sub43 = doc.add_paragraph()
    sub43.paragraph_format.space_before = Pt(10)
    sub43.paragraph_format.space_after = Pt(3)
    sub43.paragraph_format.keep_with_next = True
    add_runs(sub43, "4.3 Studi Ablasi Fitur Bertahap (Feature Ablation Study)", default_size=Pt(11), default_bold=True, default_italic=True)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.35)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    add_runs(p, "Tabel 4 dan Gambar 5 menyajikan hasil evaluasi ablasi fitur beserta analisis interval kepercayaan bootstrap.", default_size=Pt(11))

    pt4 = doc.add_paragraph()
    pt4.paragraph_format.space_before = Pt(8)
    pt4.paragraph_format.space_after = Pt(2)
    add_runs(pt4, "**Tabel 4. Hasil Studi Ablasi Fitur Menggunakan Model XGBoost**", default_size=Pt(10))

    pnot4 = doc.add_paragraph()
    pnot4.paragraph_format.space_after = Pt(4)
    add_runs(pnot4, "*(Catatan metodologis: Evaluasi ablasi dijalankan dengan retraining model XGBoost secara terpisah pada setiap subset fitur menggunakan regularisasi default tanpa parameter subsample/colsample khusus pada Tabel 3, menghasilkan ROC-AUC model penuh sebesar 0,7684 yang selaras secara operasional dengan 0,7682 pada model utama Tabel 3).* ", default_size=Pt(8.5), default_italic=True)

    t4_headers = ["Konfigurasi Fitur", "ROC-AUC", "PR-AUC (vs Base 0,25)", "Top-20% Risk Yield [95% CI]", "Cumulative Lift", "Temuan Empiris"]
    t4_rows = [
        ["**1. Data SPT Mandiri Saja**", "0,5641", "0,3192", "27,2% [22,8%, 31,6%]", "1,36x", "Pelaporan mandiri semata memiliki daya pembeda yang sangat terbatas"],
        ["**2. + Transaksi Digital (Gateway & GMV)**", "0,7348", "0,5197", "43,6% [38,8%, 48,4%]", "2,18x", "Peningkatan substansial (+0,1707 AUC); menangkap volume peredaran usaha"],
        ["**3. + Data Logistik Pengiriman**", "0,7667", "0,5455", "46,0% [41,2%, 50,4%]", "2,30x", "Peningkatan diskriminasi tertinggi melalui verifikasi pergerakan barang fisik"],
        ["**4. Model Penuh (+ 2 Indikator Makro BPS)**", "**0,7684**", "**0,5467**", "**45,2% [40,8%, 49,6%]**", "**2,26x**", "Memberikan konteks makro wilayah (marginal ΔAUC = +0,0017)"]
    ]
    t4_widths = [1.8, 0.7, 0.8, 1.2, 0.6, 1.2]
    t4_aligns = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]
    populate_table(doc, t4_headers, t4_rows, t4_widths, t4_aligns, font_size=Pt(9.0))

    if os.path.exists("images/figure5_ablation_study_bars.png"):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(2)
        r_img = p_img.add_run()
        r_img.add_picture("images/figure5_ablation_study_bars.png", width=Inches(5.6))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(8)
        add_runs(p_cap, "*Gambar 5. Perbandingan Skor ROC-AUC pada Studi Ablasi Fitur (300 DPI).*", default_size=Pt(9.5))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.35)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    add_runs(p, "**Analisis Penyelarasan Top-20% Yield dan ROC-AUC:** Berdasarkan Tabel 4, penambahan indikator makro regional BPS meningkatkan ROC-AUC dari 0,7667 menjadi 0,7684, namun nilai titik (point estimate) Top-20% Yield mengalami pergeseran marjinal dari 46,0% ke 45,2%. Analisis bootstrap (B=500) membuktikan bahwa interval kepercayaan 95% antara Tahap 3 ([41,2%, 50,4%]) dan Tahap 4 ([40,8%, 49,6%]) saling beririsan secara luas (broadly overlapping). Secara empiris, penambahan fitur makro menyebabkan sedikit penataan ulang (re-ordering) probabilitas prediksi pada observasi di sekitar batas ambang desil kedua, namun tidak mencerminkan penurunan performa pemeringkatan yang signifikan secara statistik. Di samping itu, nilai PR-AUC pada konfigurasi SPT mandiri (0,3192) hanya mencatatkan keunggulan tipis di atas garis dasar prevalensi 0,2500 (rasio 1,28x), mengonfirmasi bahwa pelaporan mandiri semata hampir tidak mengandung sinyal prediktif diskriminatif yang memadai bagi fiskus.", default_size=Pt(11))

    sub44 = doc.add_paragraph()
    sub44.paragraph_format.space_before = Pt(10)
    sub44.paragraph_format.space_after = Pt(3)
    sub44.paragraph_format.keep_with_next = True
    add_runs(sub44, "4.4 Atribusi Prediktif Berbasis SHAP sebagai Sanity Check Keselarasan DGP", default_size=Pt(11), default_bold=True, default_italic=True)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.35)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    add_runs(p, "Gambar 6 menampilkan distribusi nilai SHAP global pada model XGBoost.", default_size=Pt(11))

    if os.path.exists("images/figure6_shap_beeswarm.png"):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(2)
        r_img = p_img.add_run()
        r_img.add_picture("images/figure6_shap_beeswarm.png", width=Inches(5.6))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(8)
        add_runs(p_cap, "*Gambar 6. Distribusi Nilai SHAP: Kontribusi Fitur terhadap Prediksi Risiko (300 DPI).*", default_size=Pt(9.5))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.35)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    add_runs(p, "Analisis SHAP dalam studi simulasi ini berfungsi sebagai sanity check untuk memverifikasi bahwa model mengekstrak sinyal yang selaras dengan mekanisme DGP (verification of DGP alignment). Berdasarkan metrik mean absolute SHAP value (mean(|SHAP|)), variabel rasio pembayaran digital (PayRatio) dan keterlacakan logistik (Logistics) diidentifikasi sebagai variabel dengan atribusi prediktif tertinggi. Hal ini mengonfirmasi bahwa model secara konsisten memanfaatkan proksi yang paling sensitif terhadap perilaku laten δ_{cash} dan δ_{log}, tanpa mengklaim adanya mekanisme kausalitas langsung di lapangan.", default_size=Pt(11))

    sub45 = doc.add_paragraph()
    sub45.paragraph_format.space_before = Pt(10)
    sub45.paragraph_format.space_after = Pt(3)
    sub45.paragraph_format.keep_with_next = True
    add_runs(sub45, "4.5 Evaluasi Matriks Konfusi", default_size=Pt(11), default_bold=True, default_italic=True)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.35)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    add_runs(p, "Gambar 7 menampilkan matriks konfusi ternormalisasi pada data uji independen.", default_size=Pt(11))

    if os.path.exists("images/figure7_confusion_matrix.png"):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(2)
        r_img = p_img.add_run()
        r_img.add_picture("images/figure7_confusion_matrix.png", width=Inches(5.0))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(8)
        add_runs(p_cap, "*Gambar 7. Normalized Confusion Matrix Model XGBoost (300 DPI).*", default_size=Pt(9.5))

    sub46 = doc.add_paragraph()
    sub46.paragraph_format.space_before = Pt(10)
    sub46.paragraph_format.space_after = Pt(3)
    sub46.paragraph_format.keep_with_next = True
    add_runs(sub46, "4.6 Uji Generalisasi Lintas Provinsi dan Sensitivitas Bobot DGP", default_size=Pt(11), default_bold=True, default_italic=True)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.35)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    add_runs(p, "Tabel 5 dan Gambar 8 merangkum hasil uji sensitivitas model terhadap variasi bobot parameter laten serta validasi lintas provinsi.", default_size=Pt(11))

    pt5 = doc.add_paragraph()
    pt5.paragraph_format.space_before = Pt(8)
    pt5.paragraph_format.space_after = Pt(2)
    add_runs(pt5, "**Tabel 5. Hasil Uji Sensitivitas Spesifikasi Parameter Bobot Laten DGP**", default_size=Pt(10))

    pnot5 = doc.add_paragraph()
    pnot5.paragraph_format.space_after = Pt(4)
    add_runs(pnot5, "*(Catatan metodologis: Setiap skenario sensitivitas dibangkitkan sebagai eksperimen replikasi independen berbasis generator random state lokal; oleh karena itu, nilai ROC-AUC Skenario A [0,7804] tidak dimaksudkan sebagai replikasi numerik identik terhadap dataset utama Tabel 3 [0,7856]).*", default_size=Pt(8.5), default_italic=True)

    t5_headers = ["Skenario Bobot Laten DGP (δ_cash / δ_inv / δ_log)", "Karakteristik Skenario", "Logistic Regression ROC-AUC", "XGBoost ROC-AUC"]
    t5_rows = [
        ["**Skenario A (0,40 / 0,35 / 0,25)**", "Baseline Seimbang", "0,7804", "0,7531"],
        ["**Skenario B (0,25 / 0,50 / 0,25)**", "Dominan Distorsi Inventaris", "0,7430", "0,6931"],
        ["**Skenario C (0,30 / 0,25 / 0,45)**", "Dominan Anomali Logistik", "0,7681", "0,7648"]
    ]
    t5_widths = [2.2, 1.8, 1.1, 1.2]
    t5_aligns = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER]
    populate_table(doc, t5_headers, t5_rows, t5_widths, t5_aligns, font_size=Pt(9.5))

    if os.path.exists("images/figure8_geographical_holdout.png"):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(2)
        r_img = p_img.add_run()
        r_img.add_picture("images/figure8_geographical_holdout.png", width=Inches(5.6))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(8)
        add_runs(p_cap, "*Gambar 8. Evaluasi Validasi Spasial Lintas R=5 Pasangan Provinsi Holdout terhadap Garis Dasar Acak (0,50) dan Model Uji Penuh (300 DPI). Sumbu Y menunjukkan skor ROC-AUC pada kelompok data holdout provinsi yang tidak terlihat dalam proses pelatihan.*", default_size=Pt(9.5))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.35)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    add_runs(p, "Validasi lintas provinsi berulang (R=5 pasangan provinsi holdout) pada Gambar 8 menghasilkan skor rata-rata **0,7643 ± 0,0321**. Seluruh skor holdout provinsi berada jauh di atas garis dasar tebakan acak (0,50) dan mendekati performa model uji penuh (0,7682). Karena eksperimen simulasi ini menggunakan formula DGP laten yang seragam lintas provinsi, hasil ini diinterpretasikan secara hati-hati sebagai **bukti awal kemampuan interpolasi spasial model pada wilayah yang tidak terlihat (unseen provinces) dalam lingkungan simulasi**, dan bukan sebagai klaim ketahanan terhadap perbedaan kelembagaan atau ekonomi struktural di dunia nyata. Hasil Tabel 5 menunjukkan bahwa performa model mengalami variasi (XGBoost turun ke 0,6931 pada Skenario B di mana distorsi inventaris dominan), yang mencerminkan sensitivitas model terhadap perubahan struktur risiko laten yang diasumsikan.", default_size=Pt(11))

    # 5. PEMBAHASAN DAN IMPLIKASI TATA KELOLA
    h5 = doc.add_paragraph()
    h5.paragraph_format.space_before = Pt(14)
    h5.paragraph_format.space_after = Pt(4)
    h5.paragraph_format.keep_with_next = True
    add_runs(h5, "5. PEMBAHASAN DAN IMPLIKASI TATA KELOLA", default_size=Pt(12), default_bold=True)
    
    sub51 = doc.add_paragraph()
    sub51.paragraph_format.space_before = Pt(10)
    sub51.paragraph_format.space_after = Pt(3)
    sub51.paragraph_format.keep_with_next = True
    add_runs(sub51, "5.1 Implikasi Regulasi, Tata Kelola Data, dan Kepatuhan UU PDP", default_size=Pt(11), default_bold=True, default_italic=True)
    
    paragraphs_51 = [
        "Penerapan praktis dari kerangka kerja pemeringkatan risiko berbasis data pihak ketiga menuntut kepatuhan ketat terhadap kerangka hukum yang berlaku di Indonesia:",
        "a. **Dasar Hukum Integrasi Data Pihak Ketiga:** Integrasi data transaksi perbankan dan lokapasar memerlukan landasan hukum sektoral perpajakan (seperti mandat pelaporan data instansi, lembaga, asosiasi, dan pihak lain/ILAP berdasarkan UU KUP dan PMK terkait) yang diselaraskan dengan asas pemrosesan data spesifik pada UU Nomor 27 Tahun 2022 (UU PDP).",
        "b. **Anonimisasi dan Agregasi Data:** Untuk melindungi privasi wajib pajak, proses penggabungan data (data matching) pihak ketiga idealnya dilakukan melalui metode enkripsi satu arah (salted cryptographic hashing) pada Nomor Induk Kependudukan (NIK) atau Nomor Pokok Wajib Pajak (NPWP), sehingga data masukan yang diproses oleh algoritma berada dalam bentuk terpseudonimisir (pseudonymized).",
        "c. **Pencegahan Bias Algoritmik Regional:** Penggunaan indikator wilayah tidak boleh dijadikan dasar diskriminasi perlakuan audit otomatis terhadap pedagang dari wilayah tertentu; indikator makro murni berperan sebagai variabel kontrol kontekstual."
    ]
    for para in paragraphs_51:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if para.startswith("a.") or para.startswith("b.") or para.startswith("c."):
            p.paragraph_format.left_indent = Inches(0.25)
        else:
            p.paragraph_format.first_line_indent = Inches(0.35)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        add_runs(p, para, default_size=Pt(11))

    sub52 = doc.add_paragraph()
    sub52.paragraph_format.space_before = Pt(10)
    sub52.paragraph_format.space_after = Pt(3)
    sub52.paragraph_format.keep_with_next = True
    add_runs(sub52, "5.2 Pertimbangan Ambang Batas Operasional dan Kompromi Kesalahan Audit", default_size=Pt(11), default_bold=True, default_italic=True)
    
    paragraphs_52 = [
        "Dalam implementasi praktis di otoritas perpajakan, pemilihan wajib pajak yang diaudit tidak didasarkan pada ambang batas probabilitas arbitrer (seperti 0,50), melainkan ditentukan oleh **kapasitas audit operasional (audit capacity)** dan pertimbangan rasio biaya kesalahan:",
        "a. **Biaya Sosial Positif Palsu (False Positive Cost):** Pemanggilan pemeriksaan terhadap wajib pajak yang sebenarnya patuh menimbulkan beban biaya kepatuhan (compliance cost), hilangnya waktu produktif pelaku usaha UMKM, serta risiko mengikis kepercayaan publik (tax morale) terhadap asas keadilan fiskal (Alm & Malézieux, 2021).",
        "b. **Biaya Fiskal Negatif Palsu (False Negative Cost):** Lolosnya wajib pajak tidak patuh dari pengawasan menyebabkan potensi kebocoran penerimaan negara (tax gap) dan menciptakan ketidakadilan persaingan usaha (unfair competition) terhadap pedagang yang patuh membayar pajak.",
        "c. **Penetapan Kuota Pemeriksaan Berbasis Persentil (Capacity-Agnostic Top-K% Selection):** Mengingat keterbatasan personel pemeriksa pajak, DJP dapat memanfaatkan model murni untuk memeringkat risiko wajib pajak secara kontinu, kemudian memilih kuota K% teratas (misalnya 10% atau 20% teratas) sesuai kapasitas sumber daya unit vertikal kantor pelayanan pajak setempat."
    ]
    for para in paragraphs_52:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if para.startswith("a.") or para.startswith("b.") or para.startswith("c."):
            p.paragraph_format.left_indent = Inches(0.25)
        else:
            p.paragraph_format.first_line_indent = Inches(0.35)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        add_runs(p, para, default_size=Pt(11))

    sub53 = doc.add_paragraph()
    sub53.paragraph_format.space_before = Pt(10)
    sub53.paragraph_format.space_after = Pt(3)
    sub53.paragraph_format.keep_with_next = True
    add_runs(sub53, "5.3 Keterbatasan Simulasi dan Validasi Eksternal", default_size=Pt(11), default_bold=True, default_italic=True)
    
    paragraphs_53 = [
        "Penelitian ini memiliki sejumlah keterbatasan epistemik dan metodologis yang harus ditekankan secara transparan:",
        "a. **Karakteristik Proof-of-Concept:** Seluruh temuan kuantitatif (seperti ROC-AUC 0,7856 dan Decile Lift 2,38x) diperoleh dari lingkungan simulasi sintetis berbasis asumsi DGP yang dirancang oleh penulis. Hasil ini berfungsi sebagai pembuktian kelayakan metodologis (proof-of-concept) dan belum dapat diekstrapolasi langsung sebagai evaluasi efektivitas kebijakan empiris pada populasi wajib pajak riil (Snoke et al., 2018).",
        "b. **Keterbatasan Dua Indikator BPS:** Evaluasi makroekonomi dalam penelitian ini hanya mencakup dua indikator (penetrasi e-commerce dan IP-TIK). Kesimpulan bahwa data makro memberikan kontribusi marginal hanya berlaku khusus pada dua indikator yang diuji dalam DGP ini, dan tidak dapat digeneralisasikan pada seluruh portofolio data statistik BPS.",
        "c. **Kebutuhan Validasi Lapangan:** Validasi eksternal memerlukan akses terhadap sampel audit historis teranomisasi (historical tax audit sample) dari otoritas pajak untuk menguji apakah korelasi antara proksi pihak ketiga dan temuan audit riil memiliki kekuatan sinyal yang setara dengan asumsi DGP."
    ]
    for para in paragraphs_53:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if para.startswith("a.") or para.startswith("b.") or para.startswith("c."):
            p.paragraph_format.left_indent = Inches(0.25)
        else:
            p.paragraph_format.first_line_indent = Inches(0.35)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        add_runs(p, para, default_size=Pt(11))

    # 6. KESIMPULAN DAN REKOMENDASI
    h6 = doc.add_paragraph()
    h6.paragraph_format.space_before = Pt(14)
    h6.paragraph_format.space_after = Pt(4)
    h6.paragraph_format.keep_with_next = True
    add_runs(h6, "6. KESIMPULAN DAN REKOMENDASI", default_size=Pt(12), default_bold=True)
    
    sub61 = doc.add_paragraph()
    sub61.paragraph_format.space_before = Pt(10)
    sub61.paragraph_format.space_after = Pt(3)
    sub61.paragraph_format.keep_with_next = True
    add_runs(sub61, "6.1 Kesimpulan", default_size=Pt(11), default_bold=True, default_italic=True)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.35)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    add_runs(p, "Penelitian ini merumuskan kerangka tolok ukur simulasi untuk menguji nilai informasi integrasi data pihak ketiga dalam pemeringkatan risiko kepatuhan pedagang daring. Dalam batas lingkungan simulasi yang diasumsikan, studi ablasi mengindikasikan bahwa pelaporan mandiri semata memiliki daya pembeda yang sangat terbatas (ROC-AUC 0,5641), sedangkan integrasi data transaksi gerbang pembayaran digital dan logistik berkorelasi dengan peningkatan performa pemeringkatan risiko yang substansial (ROC-AUC 0,7667 dan Top-20% Risk Yield 46,0%). Dua indikator regional berbasis data BPS yang diuji memberikan kontribusi prediktif marginal (ΔAUC = +0,0017). Model linier Logistic Regression memberikan kinerja diskriminasi tertinggi pada data uji (0,7856) karena kompleksitas model yang rendah dan minimnya risiko overfitting pada struktur data linier terdistorsi, sementara TPE-optimized XGBoost menghasilkan ROC-AUC 0,7682. Validasi lintas provinsi memberikan indikasi awal kemampuan generalisasi model pada provinsi holdout di dalam lingkungan simulasi (0,7643 ± 0,0321).", default_size=Pt(11))

    sub62 = doc.add_paragraph()
    sub62.paragraph_format.space_before = Pt(10)
    sub62.paragraph_format.space_after = Pt(3)
    sub62.paragraph_format.keep_with_next = True
    add_runs(sub62, "6.2 Rekomendasi Kebijakan", default_size=Pt(11), default_bold=True, default_italic=True)
    
    paragraphs_62 = [
        "a. **Interoperabilitas Data Pihak Ketiga:** Memprioritaskan standardisasi protokol pertukaran data terenkripsi antara otoritas pajak, penyedia gerbang pembayaran digital, dan platform logistik.",
        "b. **Pemanfaatan Model sebagai Alat Pemeringkat Prioritas (Decision Support System):** Menggunakan skor probabilitas model murni sebagai instrumen penyaring prioritas pemeriksaan, bukan penentu sanksi hukum otomatis, guna menjaga akuntabilitas algoritma.",
        "c. **Uji Validasi Empiris Bertahap:** Menjadikan kerangka simulasi ini sebagai landasan desain eksperimen sebelum menguji coba model pada basis data audit administratif riil berskala percontohan (pilot project)."
    ]
    for para in paragraphs_62:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        add_runs(p, para, default_size=Pt(11))

    # KETERSEDIAAN DATA & KODE
    h_avail = doc.add_paragraph()
    h_avail.paragraph_format.space_before = Pt(14)
    h_avail.paragraph_format.space_after = Pt(4)
    h_avail.paragraph_format.keep_with_next = True
    add_runs(h_avail, "PERNYATAAN KETERSEDIAAN DATA DAN KODE (DATA AND CODE AVAILABILITY)", default_size=Pt(11), default_bold=True)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.35)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    add_runs(p, "Dataset simulasi sintetis, skrip pembangkitan data (Data Generating Process), pipeline pelatihan machine learning, dan buku kerja analisis interaktif (Jupyter Notebook) tersedia secara terbuka untuk mendukung prinsip sains terbuka dan reprodusibilitas komputasional penuh melalui repositori publik: https://github.com/izamrosiawan/tax-compliance-automl.", default_size=Pt(11))

    # DAFTAR PUSTAKA
    href = doc.add_paragraph()
    href.paragraph_format.space_before = Pt(14)
    href.paragraph_format.space_after = Pt(4)
    href.paragraph_format.keep_with_next = True
    add_runs(href, "DAFTAR PUSTAKA", default_size=Pt(12), default_bold=True)
    
    references = [
        "Alm, J., & Malézieux, A. (2021). 40 years of tax evasion games: a meta-analysis. *Experimental Economics*, 24(3), 699-750. https://doi.org/10.1007/s10683-020-09679-3",
        "Badan Pusat Statistik. (2024). *Indeks Pembangunan Teknologi Informasi dan Komunikasi 2023*. Jakarta: BPS RI.",
        "Badan Pusat Statistik. (2024). *Statistik E-Commerce 2024*. Jakarta: BPS RI.",
        "Battaglini, M., Guiso, L., Lacava, C., Miller, D. L., & Patacchini, E. (2022). *Refining Public Policies with Machine Learning: The Case of Tax Auditing* (NBER Working Paper No. 30777). National Bureau of Economic Research. https://doi.org/10.3386/w30777",
        "Bergstra, J., Bardenet, R., Bengio, Y., & Kégl, B. (2011). Algorithms for hyper-parameter optimization. *Advances in Neural Information Processing Systems (NeurIPS)*, 24, 2546-2554.",
        "Carpenter, J., & Bithell, J. (2000). Bootstrap confidence intervals: when, which, what? A practical guide for medical statisticians. *Statistics in Medicine*, 19(9), 1141-1164. https://doi.org/10.1002/(sici)1097-0258(20000515)19:9<1141::aid-sim479>3.0.co;2-f",
        "de Roux, D., Pérez, B., Moreno, A., Villamil, M. P., & Figueroa, C. (2018). Tax fraud detection for under-reporting declarations using an unsupervised machine learning approach. In *Proceedings of the 24th ACM SIGKDD International Conference on Knowledge Discovery & Data Mining (KDD '18)* (pp. 215-222). Association for Computing Machinery. https://doi.org/10.1145/3219819.3219878",
        "Direktorat Jenderal Pajak. (2023). *Laporan Tahunan Direktorat Jenderal Pajak 2023: Transformasi Digital Perpajakan*. Jakarta: Kementerian Keuangan Republik Indonesia.",
        "Direktorat Jenderal Pajak. (2026). *Penerimaan Pajak Digital Capai Rp38,7 Triliun per Kuartal I-2026* (Siaran Pers No. SP-08/2026). Jakarta: Kementerian Keuangan Republik Indonesia.",
        "Feurer, M., Klein, A., Eggensperger, K., Springenberg, J., Blum, M., & Hutter, F. (2019). Auto-sklearn: Efficient and robust automated machine learning. *Automated Machine Learning*, 113-134. Springer, Cham. https://doi.org/10.1007/978-3-030-05318-5_6",
        "Google, Temasek, & Bain & Company. (2025). *e-Conomy SEA 2025: Navigating the Digital Acceleration in Southeast Asia*. Singapore.",
        "Hutter, F., Kotthoff, L., & Vanschoren, J. (Eds.). (2019). *Automated Machine Learning: Methods, Systems, Challenges*. Springer Nature. https://doi.org/10.1007/978-3-030-05318-5",
        "Kementerian Keuangan Republik Indonesia. (2024). *Kerangka Ekonomi Makro dan Pokok-Pokok Kebijakan Fiskal Tahun 2025*. Jakarta: Badan Kebijakan Fiskal.",
        "Khwaja, M. S., Awasthi, R., & Loeprick, J. (Eds.). (2011). *Risk-Based Tax Audits: Approaches and Country Experiences*. Washington, DC: The World Bank. https://doi.org/10.1596/978-0-8213-8754-2",
        "Kim, S., Tsai, Y. C., Singh, K., Choi, Y., Ibok, E., Li, C. T., & Cha, M. (2020). DATE: Dual attentive tree-aware embedding for customs fraud detection. In *Proceedings of the 26th ACM SIGKDD International Conference on Knowledge Discovery & Data Mining (KDD '20)* (pp. 2880-2890). Association for Computing Machinery. https://doi.org/10.1145/3394486.3403339",
        "Kleven, H. J., Knudsen, M. B., Kreiner, C. T., Pedersen, S., & Saez, E. (2011). Unwilling or unable to cheat? Evidence from a tax audit experiment in Denmark. *Econometrica*, 79(3), 651-692. https://doi.org/10.3982/ECTA9113",
        "Lundberg, S. M., & Lee, S. I. (2017). A unified approach to interpreting model predictions. *Advances in Neural Information Processing Systems (NeurIPS)*, 30, 4765-4774.",
        "Mascagni, N., & Mengistu, A. T. (2019). The data revolution in tax administration: Applications, opportunities and challenges. *ICTD Working Paper*, 96, 1-38.",
        "Naritomi, J. (2019). Consumers as tax auditors: Electronic invoice programs and tax compliance in Brazil. *American Economic Review*, 109(5), 1730-1772. https://doi.org/10.1257/aer.20160658",
        "Nowok, B., Raab, G. M., & Dibben, C. (2016). synthpop: Bespoke creation of synthetic data in R. *Journal of Statistical Software*, 74(11), 1-26. https://doi.org/10.18637/jss.v074.i11",
        "OECD. (2020). *Tax Challenges Arising from Digitalisation – Report on Pillar One Blueprint*. Paris: OECD Publishing. https://doi.org/10.1787/beba0634-en",
        "Perez-Truglia, R. (2020). The effects of income transparency: Evidence from digital disclosure in Norway. *Journal of Political Economy*, 128(7), 2677-2716. https://doi.org/10.1086/706798",
        "Republik Indonesia. (2022). *Undang-Undang Republik Indonesia Nomor 27 Tahun 2022 tentang Pelindungan Data Pribadi*. Lembaran Negara Republik Indonesia Tahun 2022 Nomor 196. Jakarta.",
        "Slemrod, J. (2019). Tax compliance and enforcement. *Journal of Economic Literature*, 57(4), 904-954. https://doi.org/10.1257/jel.20181437",
        "Snoke, J., Raab, G. M., Nowok, B., Dibben, C., & Slavkovic, A. (2018). General and specific utility measures for synthetic data. *Journal of the Royal Statistical Society: Series A (Statistics in Society)*, 181(3), 663-688. https://doi.org/10.1111/rssa.12358"
    ]
    for ref in references:
        pref = doc.add_paragraph()
        pref.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pref.paragraph_format.left_indent = Inches(0.3)
        pref.paragraph_format.first_line_indent = Inches(-0.3)
        pref.paragraph_format.space_after = Pt(4)
        pref.paragraph_format.line_spacing = 1.05
        add_runs(pref, ref, default_size=Pt(10))

    # LAMPIRAN A
    h_app = doc.add_paragraph()
    h_app.paragraph_format.space_before = Pt(14)
    h_app.paragraph_format.space_after = Pt(4)
    h_app.paragraph_format.keep_with_next = True
    add_runs(h_app, "LAMPIRAN A. NILAI JANGKAR INDIKATOR MAKRO PROVINSI BPS (DATA PROVENANCE)", default_size=Pt(11), default_bold=True)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.35)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    add_runs(p, "Tabel A1 merangkum nilai jangkar makro agregat 10 provinsi yang digunakan dalam proses pembangkitan data sintetis (Data Generating Process).", default_size=Pt(11))

    pta1 = doc.add_paragraph()
    pta1.paragraph_format.space_before = Pt(8)
    pta1.paragraph_format.space_after = Pt(2)
    add_runs(pta1, "**Tabel A1. Nilai Jangkar Indikator Regional BPS pada 10 Provinsi Utama**", default_size=Pt(10))

    ta1_headers = ["No", "Provinsi", "Persentase Usaha E-Commerce (%)¹", "Indeks Pembangunan TIK (IP-TIK 2023, Skala 0–10)²"]
    ta1_rows = [
        ["1", "DKI Jakarta", "63,54", "7,73"],
        ["2", "Bali", "48,74", "6,60"],
        ["3", "Banten", "44,12", "6,38"],
        ["4", "Jawa Barat", "43,38", "6,15"],
        ["5", "Jawa Timur", "36,85", "5,96"],
        ["6", "Jawa Tengah", "33,20", "5,86"],
        ["7", "Sumatera Utara", "28,41", "6,04"],
        ["8", "Sulawesi Selatan", "27,90", "6,01"],
        ["9", "Riau", "26,15", "6,07"],
        ["10", "Sumatera Selatan", "24,32", "5,88"]
    ]
    ta1_widths = [0.6, 2.1, 1.8, 1.8]
    ta1_aligns = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER]
    populate_table(doc, ta1_headers, ta1_rows, ta1_widths, ta1_aligns, font_size=Pt(9.5))

    p_src = doc.add_paragraph()
    p_src.paragraph_format.space_after = Pt(10)
    add_runs(p_src, "*Sumber data:* \n¹ Badan Pusat Statistik. (2024). *Statistik E-Commerce 2024*, Tabel 3.1.\n² Badan Pusat Statistik. (2024). *Indeks Pembangunan Teknologi Informasi dan Komunikasi 2023*, Tabel 4.", default_size=Pt(9.0), default_italic=True)

    output_filename = "tax_compliance_manuscript_sinta2_ready.docx"
    doc.save(output_filename)
    print(f"Masterpiece manuscript created successfully at: {output_filename}")

if __name__ == "__main__":
    build_perfect_manuscript()
